#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
会议录音转写助手 v3.1.0 (Web Edition)
架构：Python后端(FastAPI+WebSocket) + Edge浏览器前端
功能：系统音频/麦克风录音、FunASR/讯飞实时转写、AI纪要
"""
import sys
import os
import re
import json
import time
import wave
import queue
import threading
import subprocess
import traceback
from pathlib import Path
from datetime import datetime

# 拼音级热词纠正依赖 pypinyin + rapidfuzz，缺失时自动降级为仅解码期热词偏置
try:
    from pypinyin import lazy_pinyin
    from rapidfuzz import fuzz as _rf_fuzz
    _POSTPROCESS_HOTWORDS_AVAILABLE = True
except ImportError:
    lazy_pinyin = None
    _rf_fuzz = None
    _POSTPROCESS_HOTWORDS_AVAILABLE = False

# ─── 配置 ───
APP_NAME = "会议录音转写助手"
APP_VERSION = "3.1.0"
BASE_DIR = Path(__file__).parent
UI_DIR = BASE_DIR / "ui"
DATA_DIR = Path.home() / "MeetingRecorder"
RECORDINGS_DIR = DATA_DIR / "recordings"
TRANSCRIPTS_DIR = DATA_DIR / "transcripts"

# ─── 确保目录存在 ───
def ensure_dirs():
    for d in [DATA_DIR, RECORDINGS_DIR, TRANSCRIPTS_DIR]:
        d.mkdir(parents=True, exist_ok=True)

# ─── 输出路径安全 ───
# 可保存的转写格式，与前端 saveFormat 下拉保持一致。
# fmt 直接拼进文件后缀，不限定白名单就能落盘 .bat / .ps1 / .html 等
# 可执行或可被双击打开的文件。
ALLOWED_SAVE_FORMATS = ("txt", "docx", "srt", "vtt", "json")

# 文件名字符白名单：中文、大小写字母、数字、点、下划线、连字符。
_UNSAFE_FILENAME_CHARS = re.compile(r"[^0-9A-Za-z\u4e00-\u9fff._-]")
_MAX_FILENAME_LEN = 120
# Windows 盘符前缀（仅开头单字母 + 冒号，如 "C:"）
_DRIVE_PREFIX = re.compile(r"^[A-Za-z]:")
# Windows 保留设备名（带任意扩展名也不可用）
_WIN_RESERVED_NAMES = frozenset(
    ["CON", "PRN", "AUX", "NUL"]
    + [f"COM{i}" for i in range(1, 10)]
    + [f"LPT{i}" for i in range(1, 10)]
)


def _sanitize_filename(filename, fallback="transcript"):
    """把前端传来的文件名压成纯文件名，杜绝路径穿越。

    前端传什么都不能影响落盘位置：目录成分、盘符、`..` 全部丢掉。
    """
    name = str(filename or "").strip()
    # 只取最后一段，丢掉任何目录成分（同时兼容 / 与 \）
    name = name.replace("\\", "/").split("/")[-1]
    # 剔掉盘符前缀（"C:evil" 这种相对盘符引用）。
    # 注意不能用 split(":")[-1]：那会把所有冒号都当盘符处理，
    # "10:30会议" 会静默丢成 "30会议"。其余冒号交给字符白名单换成 _。
    name = _DRIVE_PREFIX.sub("", name)
    name = _UNSAFE_FILENAME_CHARS.sub("_", name)
    # 前导点会产生隐藏文件；"." / ".." 必须拒掉
    name = name.lstrip(".").strip("_ ")
    if not name or set(name) <= {"."}:
        name = fallback
    if len(name) > _MAX_FILENAME_LEN:
        name = name[:_MAX_FILENAME_LEN]
    if name.split(".")[0].upper() in _WIN_RESERVED_NAMES:
        name = f"_{name}"
    return name


def _resolve_output_path(base_dir, filename, ext, fallback="transcript"):
    """在 base_dir 下构造安全输出路径。

    除了净化文件名，还对 resolve 后的结果做一次归属校验——净化函数
    万一有缺口，这里仍能拦住，而不是默默写到目录外。
    """
    safe = _sanitize_filename(filename, fallback)
    base = Path(base_dir).resolve()
    path = (base / f"{safe}.{ext}").resolve()
    if path.parent != base:
        raise ValueError(f"非法输出路径: {filename!r}")
    return path

# ─── 文件转写模型 ───
# 文件转写（非实时）使用的 ASR 模型。实测在中文多说话人 / 带背景音的
# 场景下，paraformer-large-vad-punc 的同音字、数字处理、语种漂移均
# 优于 SenseVoiceSmall，且 CPU RTF 更低（约 0.12 vs 0.13），故设为
# 默认。SenseVoiceSmall 保留为可选项（多语种场景可能用得上），但不再
# 默认加载。
FILE_MODEL_PARAFORMER_LARGE = "iic/speech_paraformer-large-vad-punc_asr_nat-zh-cn-16k-common-vocab8404-pytorch"
FILE_MODEL_SENSEVOICE = "iic/SenseVoiceSmall"
_FILE_MODEL_ALIASES = {
    "paraformer-large": FILE_MODEL_PARAFORMER_LARGE,
    "sensevoice": FILE_MODEL_SENSEVOICE,
}


def _resolve_file_model(name):
    """把前端短名 / 配置值归一化为完整 FunASR 模型 ID。

    兼容历史配置：前端下拉传的是 "paraformer-large" / "sensevoice" 短名，
    旧默认值是完整 ID。未知值原样返回（假定已是完整 ID）。
    """
    if not name:
        return FILE_MODEL_PARAFORMER_LARGE
    return _FILE_MODEL_ALIASES.get(str(name).strip().lower(), name)


# ─── 默认配置 ───
DEFAULT_CONFIG = {
    "audio_source": "both",
    "engine": "FunASR",
    "funasr_model": FILE_MODEL_PARAFORMER_LARGE,
    "auto_transcribe": True,
    "xfyun_app_id": "",
    "xfyun_api_key": "",
    "xfyun_api_secret": "",
    "llm_api_key": "",
    "llm_provider": "xfyun",  # xfyun / volcengine / deepseek / qwen
    "hot_words": "",
    "domain": "通用",
    "save_format": "txt",
    "speaker_diarization": False,
    "preset_spk_num": 0,
}

# ─── 配置持久化 ───
CONFIG_FILE = DATA_DIR / "config.json"

def load_config():
    """加载配置，优先从持久化文件读取，缺失字段用默认值填充"""
    config = dict(DEFAULT_CONFIG)
    if CONFIG_FILE.exists():
        try:
            saved = json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
            config.update(saved)
        except Exception:
            pass
    return config

def save_config(config):
    """保存配置到文件"""
    try:
        CONFIG_FILE.write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass

# ─── 音频录制器 ───
# 边录边落盘的参数。
# 旧实现把整场录音的 PCM 全攒在 self._frames 里，只在 stop() 时一次性
# concatenate 后写 WAV，内存占用与会议时长成正比，且 stop() 里连锁副本
# （concatenate → *32767 → astype → tobytes）会把峰值再放大约 4 倍。
# 实测折算：8 小时混合录音 _frames 约 1.7 GB、stop() 峰值约 7 GB、
# 双路混合约 14 GB —— 会在按下停止那一刻 OOM，整场录音全丢。
# 改为录音期间由后台线程增量写 WAV，_frames 只作为短暂的交接缓冲。
#
# 两个参数按实测数据标定（每块 0.3 秒 → 3.33 块/秒）：
#   flush=0.5s  → 稳态积压仅 1.7 块，正常录音永不触发丢弃
#   上限 1200 块 → 最坏情况（双声道 48kHz，112.5 KB/块）占 132 MB，
#                 可容忍磁盘卡顿约 6 分钟
# 写盘量本身极小（16kHz 单声道 int16 = 32 KB/s），瓶颈不可能是带宽，
# 只可能是偶发卡顿，所以寘缓冲换安全很划算——录音丢一秒就是会议
# 内容永久缺失，而 132 MB 相比旧实现的 14 GB 峰值可以忽略不计。
_WRITER_FLUSH_INTERVAL_S = 0.5      # 落盘线程的轮询间隔
_WRITER_MAX_PENDING_BLOCKS = 1200   # _frames 里最多积压的块数（超过则丢最旧）

# 实时转写队列上限（块数）。每块 0.3 秒，16kHz 单声道 int16 ≈ 9.6 KB。
# 400 块 ≈ 120 秒 ≈ 3.8 MB，足够覆盖 ASR 的短时抖动；再多说明消费者
# 根本没在跑（用户只录音没开实时转写），继续堆下去纯属浪费内存。
_AUDIO_QUEUE_MAX_CHUNKS = 400


class AudioRecorder:
    """使用 sounddevice 录制音频"""
    def __init__(self):
        self.sd = None
        self.np = None
        self._recording = False
        self._frames = []
        self._stream = None
        self._stream2 = None          # 混合模式下的第二个录音流
        self._mix_buffer = {}         # 混合模式下的音频对齐缓冲
        self._mix_thread = None       # 混音线程（stop 时需 join，否则会写已关闭的文件）
        self._sample_rate = 16000
        self._channels = 1
        self._actual_sample_rate = 16000  # 实际录音采样率（可能因设备不同）
        self._actual_channels = 1         # 实际录音通道数
        self._lock = threading.Lock()
        # _frames 的交接锁：音频回调（append）与落盘线程（整体取走）之间
        # 必须原子交接。临界区只做 list 操作，磁盘 I/O 放在锁外，
        # 因为 sounddevice 回调绝不能阻塞（一阻塞就丢采样）。
        self._frames_lock = threading.Lock()
        # 实时转写队列必须有界：它的唯一消费者是
        # _realtime_transcribe_task，只录音不开实时转写时根本没人 get，
        # 无界队列会随录音时长线性增长（16kHz 单声道 int16 = 32 KB/s，
        # 8 小时约 879 MB）。
        self._audio_queue = queue.Queue(maxsize=_AUDIO_QUEUE_MAX_CHUNKS)
        self._dropped_chunks = 0
        # ── 边录边落盘 ──
        self._wav_writer = None
        self._wav_path = None
        self._wav_lock = threading.Lock()
        self._wav_samples_written = 0
        self._writer_thread = None

    def _import_deps(self):
        if self.sd is None:
            import sounddevice as sd
            import numpy as np
            self.sd = sd
            self.np = np

    def list_devices(self):
        """列出可用音频设备"""
        self._import_deps()
        devices = []
        for i, dev in enumerate(self.sd.query_devices()):
            if dev['max_input_channels'] > 0:
                devices.append({
                    "id": i,
                    "name": dev['name'],
                    "channels": dev['max_input_channels'],
                    "sample_rate": int(dev['default_samplerate']),
                })
        return devices

    def start(self, source="system", callback=None):
        """开始录音
        source: 'system' 系统音频(立体声混音), 'mic' 麦克风, 'both' 混合(系统+麦克风)
        callback: 可选，每收到一帧音频调用 callback(audio_bytes)
        """
        self._import_deps()
        if self._recording:
            return False

        self._frames = []
        self._recording = True
        self._callback = callback
        self._mix_buffer = {}
        self._dropped_chunks = 0

        # 先打开 WAV 写入器并拉起落盘线程，使录音从第一块开始就能落盘。
        # 设备参数（_actual_sample_rate / _actual_channels）在下面才确定，
        # 但重采样是在写入时读取这两个值的，而第一块数据要等流启动后
        # 才会到达，所以提前开线程是安全的。
        self._open_wav_writer()
        self._writer_thread = threading.Thread(target=self._writer_loop, daemon=True)
        self._writer_thread.start()

        # ── 查找设备 ──
        devices = self.list_devices()

        # 查找立体声混音设备
        loopback_dev = None
        for dev in devices:
            name_lower = dev['name'].lower()
            if any(kw in name_lower for kw in ['stereo mix', 'what u hear', 'loopback', 'virtual', 'vb-audio', 'voicemeeter', 'stereo input', '立体声混音', '混音']):
                loopback_dev = dev
                break

        # 查找最佳麦克风设备
        mic_dev = None
        for dev in devices:
            if '麦克风阵列' in dev['name'] and dev['sample_rate'] >= 48000 and dev['channels'] <= 4:
                mic_dev = dev
                break
        if not mic_dev:
            for dev in devices:
                if '麦克风' in dev['name'] and dev['channels'] <= 4:
                    mic_dev = dev
                    break

        # ── 根据source选择设备 ──
        if source == "both" and loopback_dev and mic_dev:
            # 混合模式：同时开立体声混音 + 麦克风
            sys_id = loopback_dev['id']
            sys_ch = min(loopback_dev['channels'], 2)
            sys_rate = int(loopback_dev['sample_rate'])
            mic_id = mic_dev['id']
            mic_ch = min(mic_dev['channels'], 2)
            mic_rate = int(mic_dev['sample_rate'])

            # 统一采样率（取两者中较高的，另一个重采样）
            target_rate = max(sys_rate, mic_rate)
            self._actual_sample_rate = target_rate
            self._actual_channels = 1  # 混合后输出单声道

            def sys_callback(indata, frames, time_info, status):
                if not self._recording: return
                mono = indata[:, 0] if indata.shape[1] > 1 else indata.flatten()
                # 重采样到target_rate
                if sys_rate != target_rate:
                    ratio = target_rate / sys_rate
                    new_len = int(len(mono) * ratio)
                    mono = self._resample_linear_f32(mono, new_len)
                with self._lock:
                    buf = self._mix_buffer.setdefault('sys', self.np.array([], dtype=self.np.float32))
                    self._mix_buffer['sys'] = self.np.concatenate([buf, mono])

            def mic_callback(indata, frames, time_info, status):
                if not self._recording: return
                mono = indata[:, 0] if indata.shape[1] > 1 else indata.flatten()
                if mic_rate != target_rate:
                    ratio = target_rate / mic_rate
                    new_len = int(len(mono) * ratio)
                    mono = self._resample_linear_f32(mono, new_len)
                with self._lock:
                    buf = self._mix_buffer.setdefault('mic', self.np.array([], dtype=self.np.float32))
                    self._mix_buffer['mic'] = self.np.concatenate([buf, mono])

            # 混音线程：定期从两个buffer取数据叠加
            import threading
            def mix_worker():
                import time as _time
                while self._recording:
                    _time.sleep(0.3)  # 每300ms混一次
                    with self._lock:
                        sys_buf = self._mix_buffer.get('sys', self.np.array([], dtype=self.np.float32))
                        mic_buf = self._mix_buffer.get('mic', self.np.array([], dtype=self.np.float32))
                        if len(sys_buf) == 0 and len(mic_buf) == 0:
                            continue
                        mixed, new_sys, new_mic = self._drain_mix_buffers(sys_buf, mic_buf)
                        self._mix_buffer['sys'] = new_sys
                        self._mix_buffer['mic'] = new_mic
                    if len(mixed) == 0:
                        continue
                    # 保存原始帧（由落盘线程增量写入 WAV，不再攒整场）
                    self._append_frame(mixed.reshape(-1, 1))
                    # 重采样到16kHz送队列
                    if target_rate != 16000:
                        ratio = 16000 / target_rate
                        new_len = int(len(mixed) * ratio)
                        mixed_16k = self._resample_linear_f32(mixed, new_len)
                    else:
                        mixed_16k = mixed
                    pcm = self._float_to_pcm16(mixed_16k)
                    self._enqueue_pcm(pcm)
                    if self._callback:
                        try: self._callback(pcm)
                        except: pass

            try:
                self._stream = self.sd.InputStream(device=sys_id, channels=sys_ch, samplerate=sys_rate, dtype='float32', blocksize=int(sys_rate*0.3), callback=sys_callback)
                self._stream2 = self.sd.InputStream(device=mic_id, channels=mic_ch, samplerate=mic_rate, dtype='float32', blocksize=int(mic_rate*0.3), callback=mic_callback)
                self._stream.start()
                self._stream2.start()
                self._mix_thread = threading.Thread(target=mix_worker, daemon=True)
                self._mix_thread.start()
                return True
            except Exception as e:
                # 混合模式失败，回退到单独麦克风
                if self._stream:
                    try: self._stream.stop(); self._stream.close()
                    except: pass
                if self._stream2:
                    try: self._stream2.stop(); self._stream2.close()
                    except: pass
                self._stream = None
                self._stream2 = None
                # 回退到麦克风模式
                if mic_dev:
                    device_id = mic_dev['id']
                    device_channels = min(mic_dev['channels'], 2)
                    device_sample_rate = int(mic_dev['sample_rate'])
                else:
                    device_id = self.sd.default.device[0]
                    device_channels = self._channels
                    device_sample_rate = self._sample_rate
                self._actual_sample_rate = device_sample_rate
                self._actual_channels = device_channels

        elif source == "mic":
            # 纯麦克风模式
            if mic_dev:
                device_id = mic_dev['id']
                device_channels = min(mic_dev['channels'], 2)
                device_sample_rate = int(mic_dev['sample_rate'])
            else:
                device_id = self.sd.default.device[0]
                device_channels = self._channels
                device_sample_rate = self._sample_rate
            self._actual_sample_rate = device_sample_rate
            self._actual_channels = device_channels

        else:
            # 系统音频模式
            if loopback_dev:
                device_id = loopback_dev['id']
                device_channels = min(loopback_dev['channels'], 2)
                device_sample_rate = int(loopback_dev['sample_rate'])
            elif mic_dev:
                device_id = mic_dev['id']
                device_channels = min(mic_dev['channels'], 2)
                device_sample_rate = int(mic_dev['sample_rate'])
            else:
                device_id = self.sd.default.device[0]
                device_channels = self._channels
                device_sample_rate = self._sample_rate
            self._actual_sample_rate = device_sample_rate
            self._actual_channels = device_channels

        # ── 单流模式（mic / system / both回退） ──
        def audio_callback(indata, frames, time_info, status):
            if not self._recording:
                return
            audio_data = indata.copy()
            self._append_frame(audio_data)
            pcm = self._resample_to_16k_mono(audio_data)
            self._enqueue_pcm(pcm)
            if self._callback:
                try:
                    self._callback(pcm)
                except:
                    pass

        try:
            self._stream = self.sd.InputStream(
                device=device_id,
                channels=device_channels,
                samplerate=device_sample_rate,
                dtype='float32',
                blocksize=int(device_sample_rate * 0.3),
                callback=audio_callback,
            )
            self._stream.start()
            return True
        except Exception as e:
            # 回退到默认设备
            try:
                device_id = self.sd.default.device[0]
                device_sample_rate = self._sample_rate
                device_channels = self._channels
                self._actual_sample_rate = device_sample_rate
                self._actual_channels = device_channels
                self._stream = self.sd.InputStream(
                    device=device_id,
                    channels=device_channels,
                    samplerate=device_sample_rate,
                    dtype='float32',
                    blocksize=int(device_sample_rate * 0.3),
                    callback=audio_callback,
                )
                self._stream.start()
                return True
            except Exception as e2:
                self._recording = False
                # 录音根本没起来，删掉刚建的空壳 WAV，不留垃圾文件
                self._discard_wav_writer()
                raise e2

    def _float_to_pcm16(self, samples):
        """float32 [-1,1] → 16bit PCM bytes

        必须先 clip：混音后幅度可能 >1.0（两路各 *0.7 相加最大 1.4），
        直接 astype(np.int16) 是**整数回绕**而非饱和，0.9+0.9 会变成 -24250，
        最响处翻转成反相尖峰（爆音），同时污染 ASR 输入。
        """
        clipped = self.np.clip(samples, -1.0, 1.0)
        return (clipped * 32767).astype(self.np.int16).tobytes()

    def _drain_mix_buffers(self, sys_buf, mic_buf):
        """混合两路缓冲，返回 (mixed, 剩余sys, 剩余mic)

        关键：**输出多少就必须消费多少**。旧实现在只有单路有数据时
        （mix_len==0）整段输出了该路，却按 mix_len==0 回写缓冲，
        等于一个采样都没消费 → 缓冲无限增长、同段音频每 300ms 重复输出
        一次（500 采样进、1500 采样出），转写结果出现大段复读。
        """
        empty = self.np.array([], dtype=self.np.float32)
        mix_len = min(len(sys_buf), len(mic_buf))
        if mix_len > 0:
            # 两路都有：对齐到最短长度叠加，各降 0.7 压低削波概率
            mixed = sys_buf[:mix_len] * 0.7 + mic_buf[:mix_len] * 0.7
            used_sys = used_mic = mix_len
        elif len(mic_buf) > 0:
            # 只有麦克风有数据（系统未播放声音时很常见）
            mixed = mic_buf
            used_sys, used_mic = 0, len(mic_buf)
        else:
            mixed = sys_buf
            used_sys, used_mic = len(sys_buf), 0
        new_sys = sys_buf[used_sys:] if used_sys < len(sys_buf) else empty
        new_mic = mic_buf[used_mic:] if used_mic < len(mic_buf) else empty
        return mixed, new_sys, new_mic

    def _resample_linear_f32(self, mono, new_len):
        """线性插值重采样，强制保持 float32。

        坑：`np.interp` 总是返回 float64（`np.linspace` 本身就是 float64），
        传进去 float32 也一样——每采样 4 字节静默变 8 字节，内存直接翻倍，
        而且会沿混音链传播（float64 + float32 仍是 float64）。
        实测回落 float32 数值无损（最大差异 0.0），因为 float32
        本就是音频采集的原始精度。
        """
        if new_len <= 0 or len(mono) == 0:
            return self.np.array([], dtype=self.np.float32)
        indices = self.np.linspace(0, len(mono) - 1, new_len)
        out = self.np.interp(indices, self.np.arange(len(mono)), mono)
        return out.astype(self.np.float32, copy=False)

    def _resample_to_16k_mono(self, audio_data):
        """将录音数据重采样为16kHz单声道16bit PCM bytes"""
        mono_16k = self._resample_to_16k_mono_float(audio_data)
        return self._float_to_pcm16(mono_16k)

    def _resample_to_16k_mono_float(self, audio_data):
        """将录音数据重采样为16kHz单声道float32 numpy array"""
        # audio_data: numpy array, shape (frames, channels) or (frames,)
        if self._actual_channels > 1:
            # 多通道 → 只取第一个通道（麦克风阵列中Ch0通常信号最强，
            # 取平均会把静默通道混入导致信号被稀释）
            mono = audio_data[:, 0]
        else:
            mono = audio_data.flatten()

        # 重采样到16kHz（简单线性插值）
        if self._actual_sample_rate != self._sample_rate:
            ratio = self._sample_rate / self._actual_sample_rate
            new_len = int(len(mono) * ratio)
            mono = self._resample_linear_f32(mono, new_len)

        return mono

    def _enqueue_pcm(self, pcm):
        """把一块 PCM 送进实时转写队列，队列满时丢最旧的一块。

        绝不能用阻塞式 put：调用方是 sounddevice 音频回调与混音线程，
        一阻塞就丢采样。队列持续满通常意味着没人在消费（用户只录音、
        没开实时转写），此时丢弃是正确选择——录音完整性由 WAV 落盘保证，
        而不是靠这个队列。
        """
        try:
            self._audio_queue.put_nowait(pcm)
            return True
        except queue.Full:
            pass
        try:
            self._audio_queue.get_nowait()      # 丢最旧的一块，腾位给新数据
            self._dropped_chunks += 1
        except queue.Empty:
            pass
        try:
            self._audio_queue.put_nowait(pcm)
            return True
        except queue.Full:
            self._dropped_chunks += 1
            return False

    # ─── 边录边落盘 ───

    def _open_wav_writer(self):
        """开启增量 WAV 写入器（16kHz 单声道 16bit，与 FunASR 输入一致）。

        旧实现在 stop() 里才建文件，录音全程数据只存内存；这里提前到
        start() 打开，使录音可以边录边写。
        """
        ensure_dirs()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = RECORDINGS_DIR / f"recording_{timestamp}.wav"
        wf = wave.open(str(filepath), 'wb')
        wf.setnchannels(1)
        wf.setsampwidth(2)              # 16bit
        wf.setframerate(self._sample_rate)
        with self._wav_lock:
            self._wav_writer = wf
            self._wav_path = filepath
            self._wav_samples_written = 0
        return filepath

    def _discard_wav_writer(self):
        """录音启动失败时收尾：关文件并删掉空壳，不留垃圾。"""
        with self._wav_lock:
            wf, path = self._wav_writer, self._wav_path
            self._wav_writer = None
            self._wav_path = None
            self._wav_samples_written = 0
        if wf is not None:
            try:
                wf.close()
            except Exception:
                pass
        if path is not None:
            try:
                Path(path).unlink()
            except Exception:
                pass

    def _take_pending_frames(self):
        """原子取走 _frames 里已积压的块。

        只在锁内做 list 交接，重采样与磁盘 I/O 都留到锁外，
        避免拖住音频回调。
        """
        with self._frames_lock:
            if not self._frames:
                return []
            blocks = self._frames
            self._frames = []
        return blocks

    def _write_blocks_to_wav(self, blocks):
        """把若干原始音频块重采样为 16kHz 单声道后追写进 WAV。

        逐块重采样而不是整段重采样：这与实时路径（audio_callback 里的
        _resample_to_16k_mono）本来就是同一做法，且线性插值在 0.3 秒分块
        边界上的差异可忽略；换来的是内存不再随时长增长。
        """
        if not blocks:
            return 0
        written = 0
        for block in blocks:
            try:
                mono_16k = self._resample_to_16k_mono_float(block)
                if len(mono_16k) == 0:
                    continue
                pcm = self._float_to_pcm16(mono_16k)
            except Exception:
                traceback.print_exc()
                continue
            with self._wav_lock:
                if self._wav_writer is None:
                    return written
                try:
                    self._wav_writer.writeframes(pcm)
                except Exception:
                    traceback.print_exc()
                    return written
                self._wav_samples_written += len(mono_16k)
            written += len(mono_16k)
        return written

    def _writer_loop(self):
        """后台落盘线程：定期把 _frames 刷进 WAV，使内存占用与时长解耦。"""
        while self._recording:
            time.sleep(_WRITER_FLUSH_INTERVAL_S)
            try:
                self._write_blocks_to_wav(self._take_pending_frames())
            except Exception:
                # 落盘线程绝不能因单次异常退出，否则 _frames 会重新无界增长
                traceback.print_exc()

    def _append_frame(self, block):
        """音频回调侧的入口：攒入交接缓冲，并保证缓冲有界。

        正常情况下落盘线程每秒会取空它，积压应远小于上限。
        若落盘线程卡住或已死，丢最旧的块保内存不爆——宁愿丢几秒，
        不能重现“按停止那一刻 OOM、整场录音全丢”。
        """
        with self._frames_lock:
            self._frames.append(block)
            over = len(self._frames) - _WRITER_MAX_PENDING_BLOCKS
            if over > 0:
                # 批量裁剪，不用 del self._frames[0]：后者是 O(n) 指针搬移，
                # 上限 1200 时满缓冲下每次 append 都要搬一遍。
                # 也不用 deque(maxlen)：它会静默丢弃，而丢数据必须留痕迹。
                del self._frames[:over]
                self._dropped_chunks += over

    def stop(self):
        """停止录音，返回WAV文件路径

        旧实现在这里做全量 `concatenate(self._frames)` 再整段重采样、转 PCM、
        `tobytes()`，一条链上同时存在多份副本，峰值约是 _frames 的 4 倍，
        8 小时混合录音会在“按下停止”那一刻 OOM，整场录音全丢。

        现在数据已由落盘线程写入，本方法只负责收尾：停流 → 等生产者退出
        → 刷剩余缓冲 → 关文件。内存占用与会议时长无关。
        """
        if not self._recording:
            return None
        # 先置位：混音线程与落盘线程都以它作为退出条件
        self._recording = False
        if self._stream:
            self._stream.stop()
            self._stream.close()
            self._stream = None
        if self._stream2:
            self._stream2.stop()
            self._stream2.close()
            self._stream2 = None

        # 必须先 join 混音线程：它也往 _frames 里写，不等它退出就可能在
        # 关文件之后才产出数据（丢尾巴）。超时取比 0.3s 循环间隔宽裕的值。
        for th in (self._mix_thread, self._writer_thread):
            if th is not None and th.is_alive():
                try:
                    th.join(timeout=3.0)
                except Exception:
                    pass
        self._mix_thread = None
        self._writer_thread = None

        # 把最后没来得及落盘的尾巴写完
        try:
            self._write_blocks_to_wav(self._take_pending_frames())
        except Exception:
            traceback.print_exc()

        with self._wav_lock:
            wf, path = self._wav_writer, self._wav_path
            total = self._wav_samples_written
            self._wav_writer = None
            self._wav_path = None
        if wf is not None:
            try:
                wf.close()
            except Exception:
                traceback.print_exc()

        self._frames = []
        if self._dropped_chunks:
            # 不静默：丢数据必须留下痕迹，否则没人知道录音缺了一段
            print(f"[recorder] 录音期间丢弃 {self._dropped_chunks} 块数据"
                  f"（落盘或消费跟不上）")

        # 一帧没录到（设备无输入/立即停止）——删掉只有头的空 WAV，
        # 保持与旧行为一致：无数据时返回 None
        if path is None or total <= 0:
            if path is not None:
                try:
                    Path(path).unlink()
                except Exception:
                    pass
            return None
        return str(path)

    def get_audio_chunk(self, timeout=0.1):
        """获取一个音频chunk（用于实时转写）"""
        try:
            return self._audio_queue.get(timeout=timeout)
        except queue.Empty:
            return None

    @property
    def is_recording(self):
        return self._recording


# ─── FunASR 实时转写 ───
class FunASRTranscriber:
    def __init__(self):
        self.file_model = None       # 文件转写模型 (默认 paraformer-large-vad-punc)
        self.stream_model = None     # 实时流式模型 (paraformer-online)
        self.punc_model = None       # 标点恢复模型 (punc_ct-transformer)
        self.diarization_model = None  # 说话人分离pipeline (paraformer-large-vad-punc + ERes2NetV2)
        self._file_model_name = None
        self._stream_model_name = None
        self._file_loading = False
        self._stream_loading = False
        self._file_ready = False
        self._stream_ready = False
        self._diarization_loading = False

    @staticmethod
    def _parse_hot_words(raw):
        """把配置里的热词字符串解析为列表。支持逗号/顿号/空格/换行分隔。"""
        if not raw:
            return []
        if isinstance(raw, (list, tuple)):
            items = [str(w).strip() for w in raw]
        else:
            items = re.split(r"[,，、;；\s\n]+", str(raw))
        seen, out = set(), []
        for w in items:
            w = w.strip()
            if w and w not in seen:
                seen.add(w)
                out.append(w)
        return out

    def build_hotword_kwargs(self, hot_words=None):
        """构造 FunASR 热词参数。

        FunASR 提供两层热词能力，这里同时启用：
        - hotword: 解码期偏置，对模型支持的场景直接提高召回
        - postprocess_hotwords: 解码后基于拼音的文本纠正，覆盖模型未偏置到的情况
          （依赖 pypinyin + rapidfuzz，缺失时自动跳过而不报错）
        """
        words = self._parse_hot_words(hot_words)
        if not words:
            return {}
        kwargs = {"hotword": " ".join(words)}
        if _POSTPROCESS_HOTWORDS_AVAILABLE:
            kwargs["postprocess_hotwords"] = words
            kwargs["postprocess_hotword_threshold"] = 0.85
        return kwargs

    def add_punctuation(self, text):
        """用标点模型给文本加标点"""
        if not self.punc_model or not text or not text.strip():
            return text
        try:
            result = self.punc_model.generate(input=text)
            if result and len(result) > 0:
                punctuated = result[0].get("text", text)
                return punctuated
            return text
        except Exception:
            return text

    def load_file_model(self, model_name=FILE_MODEL_PARAFORMER_LARGE, status_callback=None):
        """加载文件转写模型（默认 paraformer-large-vad-punc + VAD + 标点）"""
        if self._file_ready and self._file_model_name == model_name:
            return True
        if self._file_loading:
            for _ in range(240):
                time.sleep(0.5)
                if self._file_ready and self._file_model_name == model_name:
                    return True
                if not self._file_loading:
                    break
            return self._file_ready
        self._file_loading = True
        self._file_ready = False
        try:
            if status_callback:
                status_callback("loading", "导入FunASR...")
            from funasr import AutoModel
            if status_callback:
                status_callback("loading", "加载文件转写模型（首次需下载）...")
            self.file_model = AutoModel(
                model=model_name,
                vad_model="iic/speech_fsmn_vad_zh-cn-16k-common-pytorch",
                punc_model="iic/punc_ct-transformer_cn-en-common-vocab471067-large",
                disable_update=True,
            )
            self._file_model_name = model_name
            self._file_ready = True
            if status_callback:
                status_callback("ready", "文件转写模型就绪")
            return True
        except ImportError:
            self.file_model = None
            if status_callback:
                status_callback("error", "FunASR未安装，请运行: pip install funasr")
            return False
        except Exception as e:
            self.file_model = None
            if status_callback:
                status_callback("error", f"文件转写模型加载失败: {str(e)}")
            return False
        finally:
            self._file_loading = False

    def load_stream_model(self, model_name="iic/speech_paraformer-large_asr_nat-zh-cn-16k-common-vocab8404-online", status_callback=None):
        """加载实时流式模型 + 标点模型"""
        if self._stream_ready and self._stream_model_name == model_name:
            return True
        if self._stream_loading:
            for _ in range(240):
                time.sleep(0.5)
                if self._stream_ready and self._stream_model_name == model_name:
                    return True
                if not self._stream_loading:
                    break
            return self._stream_ready
        self._stream_loading = True
        self._stream_ready = False
        try:
            if status_callback:
                status_callback("loading", "加载实时转写模型...")
            from funasr import AutoModel
            # 实时流式ASR模型（不带VAD和标点，加载快）
            self.stream_model = AutoModel(
                model=model_name,
                disable_update=True,
            )
            # 独立加载标点模型，用于实时转写文本的后处理
            if status_callback:
                status_callback("loading", "加载标点模型...")
            self.punc_model = AutoModel(
                model="iic/punc_ct-transformer_cn-en-common-vocab471067-large",
                disable_update=True,
            )
            self._stream_model_name = model_name
            self._stream_ready = True
            if status_callback:
                status_callback("ready", "实时转写模型就绪")
            return True
        except ImportError:
            self.stream_model = None
            if status_callback:
                status_callback("error", "FunASR未安装，请运行: pip install funasr")
            return False
        except Exception as e:
            self.stream_model = None
            if status_callback:
                status_callback("error", f"实时转写模型加载失败: {str(e)}")
            return False
        finally:
            self._stream_loading = False

    # 兼容旧接口
    def load_model(self, model_name=FILE_MODEL_PARAFORMER_LARGE, status_callback=None):
        """兼容接口：根据模型名自动选择加载方法"""
        if "online" in model_name or "streaming" in model_name:
            return self.load_stream_model(model_name, status_callback)
        return self.load_file_model(model_name, status_callback)

    def transcribe_file(self, filepath, status_callback=None, hot_words=None, model_name=None):
        """转写音频文件（默认 paraformer-large-vad-punc）。

        model_name 可传完整 FunASR 模型 ID 或前端短名
        （paraformer-large / sensevoice）；None 时用默认模型。
        """
        model_name = _resolve_file_model(model_name)
        if not self.load_model(model_name=model_name, status_callback=status_callback):
            return None
        try:
            if status_callback:
                status_callback("transcribing", "正在转写...")
            kwargs = dict(input=filepath, batch_size_s=300, use_itn=True)
            kwargs.update(self.build_hotword_kwargs(hot_words))
            result = self.file_model.generate(**kwargs)
            if result and len(result) > 0:
                text = result[0].get("text", "")
                # 去掉模型输出里的特殊 token（如 SenseVoice 的
                # <|zh|>/<|NEUTRAL|> 等）；paraformer 无此类 token，此处理无害。
                text = re.sub(r'<\s*\|[^>]*?\|\s*>', '', text).strip()
                return text
            return ""
        except Exception as e:
            if status_callback:
                status_callback("error", f"转写失败: {str(e)}")
            return None

    def transcribe_file_with_diarization(self, filepath, status_callback=None, preset_spk_num=None, hot_words=None):
        """转写音频文件并做说话人分离（使用paraformer-large-vad-punc + ERes2NetV2）"""
        # 加载带说话人分离的pipeline（需要支持timestamp的ASR模型）
        if not self._load_diarization_model(status_callback):
            return None
        try:
            if status_callback:
                status_callback("transcribing", "正在转写（含说话人分离）...")
            kwargs = dict(input=filepath, batch_size_s=300, return_spk_res=True)
            if preset_spk_num:
                kwargs["preset_spk_num"] = preset_spk_num
            kwargs.update(self.build_hotword_kwargs(hot_words))
            result = self.diarization_model.generate(**kwargs)
            if result and len(result) > 0:
                item = result[0]
                text = item.get("text", "")
                text = re.sub(r'<\s*\|[^>]*?\|\s*>', '', text).strip()
                sentence_info = item.get("sentence_info", [])
                # 构建结构化结果
                structured = {
                    "text": text,
                    "sentence_info": sentence_info,  # [{text, start, end, spk, timestamp}, ...]
                    "speaker_count": max([s.get("spk", 0) for s in sentence_info], default=0) + 1 if sentence_info else 0,
                }
                return structured
            return {"text": "", "sentence_info": [], "speaker_count": 0}
        except Exception as e:
            if status_callback:
                status_callback("error", f"说话人分离转写失败: {str(e)}")
            return None

    def _load_diarization_model(self, status_callback=None):
        """加载说话人分离pipeline（独立于file_model和stream_model）"""
        if self.diarization_model is not None:
            return True
        if self._diarization_loading:
            for _ in range(240):
                time.sleep(0.5)
                if self.diarization_model is not None:
                    return True
                if not self._diarization_loading:
                    break
            return self.diarization_model is not None
        self._diarization_loading = True
        try:
            if status_callback:
                status_callback("loading", "加载说话人分离模型（首次较慢）...")
            from funasr import AutoModel
            self.diarization_model = AutoModel(
                model="iic/speech_paraformer-large-vad-punc_asr_nat-zh-cn-16k-common-vocab8404-pytorch",
                spk_model="iic/speech_eres2netv2_sv_zh-cn_16k-common",
                disable_update=True,
            )
            if status_callback:
                status_callback("ready", "说话人分离模型就绪")
            return True
        except Exception as e:
            self.diarization_model = None
            if status_callback:
                status_callback("error", f"说话人分离模型加载失败: {str(e)}")
            return False
        finally:
            self._diarization_loading = False


# ─── 讯飞实时转写 ───
class XfyunTranscriber:
    """讯飞语音转写WebSocket API"""
    def __init__(self, app_id, api_key, api_secret):
        self.app_id = app_id
        self.api_key = api_key
        self.api_secret = api_secret
        self._ws = None
        self._result_text = ""
        self._connected = False

    def _create_url(self):
        """生成讯飞鉴权URL"""
        import hmac
        import base64
        from hashlib import sha1
        from urllib.parse import urlencode, quote
        import datetime as dt

        url = "wss://iat-api.xfyun.cn/v2/iat"
        now = dt.datetime.utcnow().strftime('%a, %d %b %Y %H:%M:%S GMT')
        signature_origin = f"host: iat-api.xfyun.cn\ndate: {now}\nGET /v2/iat HTTP/1.1"
        signature_sha = hmac.new(
            self.api_secret.encode('utf-8'),
            signature_origin.encode('utf-8'),
            sha1
        ).digest()
        signature = base64.b64encode(signature_sha).decode('utf-8')
        authorization = base64.b64encode(
            f'api_key="{self.api_key}", algorithm="hmac-sha1", headers="host date request-line", signature="{signature}"'.encode('utf-8')
        ).decode('utf-8')
        params = {"authorization": authorization, "date": now, "host": "iat-api.xfyun.cn"}
        return url + '?' + urlencode(params)

    def start(self, audio_callback, status_callback=None):
        """启动讯飞实时转写"""
        import websocket
        self._result_text = ""
        self._audio_callback = audio_callback
        self._status_callback = status_callback

        def on_message(ws, message):
            try:
                data = json.loads(message)
                code = data.get("header", {}).get("code")
                if code != 0:
                    if status_callback:
                        status_callback("error", f"讯飞错误: {data}")
                    return
                result = data.get("payload", {}).get("result", {})
                text = result.get("text", "")
                if text:
                    import base64
                    import zlib
                    decoded = zlib.decompress(base64.b64decode(text), 16 + zlib.MAX_WBITS).decode('utf-8')
                    ws_data = json.loads(decoded)
                    cn = ws_data.get("cn", {})
                    st = cn.get("st", {})
                    rt = st.get("rt", [])
                    seg_text = ""
                    for r in rt:
                        for w in r.get("ws", []):
                            for c in w.get("cw", []):
                                seg_text += c.get("w", "")
                    if seg_text:
                        is_end = st.get("type") == "2"
                        self._result_text += seg_text
                        if audio_callback:
                            audio_callback(seg_text, is_end)
            except Exception as e:
                if status_callback:
                    status_callback("error", f"解析讯飞结果失败: {e}")

        def on_error(ws, error):
            if status_callback:
                status_callback("error", f"讯飞连接错误: {error}")

        def on_close(ws, close_status, close_msg):
            self._connected = False

        def on_open(ws):
            self._connected = True
            if status_callback:
                status_callback("ready", "讯飞实时转写就绪")

        try:
            url = self._create_url()
            self._ws = websocket.WebSocketApp(
                url,
                on_open=on_open,
                on_message=on_message,
                on_error=on_error,
                on_close=on_close,
            )
            self._ws_thread = threading.Thread(target=self._ws.run_forever, daemon=True)
            self._ws_thread.start()
            return True
        except Exception as e:
            if status_callback:
                status_callback("error", f"讯飞连接失败: {e}")
            return False

    def send_audio(self, audio_bytes):
        """发送音频数据到讯飞"""
        if not self._ws or not self._connected:
            return
        import base64
        audio_b64 = base64.b64encode(audio_bytes).decode('utf-8')
        frame = {
            "header": {
                "app_id": self.app_id,
                "status": 1,  # 1=数据, 2=结束
            },
            "parameter": {
                "iat": {
                    "domain": "iat",
                    "language": "zh_cn",
                    "accent": "mandarin",
                    "vad_eos": 2000,
                    "result": {
                        "encoding": "utf8",
                        "compress": "raw",
                        "format": "json"
                    }
                }
            },
            "payload": {
                "audio": {
                    "encoding": "raw",
                    "sample_rate": 16000,
                    "channels": 1,
                    "bit_depth": 16,
                    "seq": int(time.time() * 1000),
                    "status": 1,
                    "data": audio_b64
                }
            }
        }
        try:
            self._ws.send(json.dumps(frame))
        except:
            pass

    def stop(self):
        """停止讯飞转写"""
        if self._ws:
            try:
                # 发送结束帧
                end_frame = {
                    "header": {"app_id": self.app_id, "status": 2},
                    "parameter": {"iat": {"domain": "iat", "language": "zh_cn", "accent": "mandarin", "vad_eos": 2000, "result": {"encoding": "utf8", "compress": "raw", "format": "json"}}},
                    "payload": {"audio": {"encoding": "raw", "sample_rate": 16000, "channels": 1, "bit_depth": 16, "seq": int(time.time() * 1000), "status": 2, "data": ""}}
                }
                self._ws.send(json.dumps(end_frame))
            except:
                pass
            try:
                self._ws.close()
            except:
                pass
        self._connected = False
        return self._result_text


# ─── 会议进程自动监控 ───
class MeetingMonitor:
    """自动检测会议软件进程，自动开始/停止录音"""
    # 常见会议软件进程名
    MEETING_PROCESSES = {
        "wemeetapp.exe": "腾讯会议",
        "wemeet.exe": "腾讯会议",
        "dingtalk.exe": "钉钉",
        "dingtalklauncher.exe": "钉钉",
        "lark.exe": "飞书",
        "bytedance-lark.exe": "飞书",
        "zoom.exe": "Zoom",
        "teams.exe": "Microsoft Teams",
        "ms-teams.exe": "Microsoft Teams",
        "voovmeeting.exe": "腾讯会议VooV",
    }

    def __init__(self):
        self._monitoring = False
        self._thread = None
        self._meeting_detected = False
        self._detected_name = ""
        self._check_interval = 3  # seconds
        self._on_start_callback = None
        self._on_stop_callback = None
        self._on_status_callback = None

    def _check_meeting_running(self):
        """检查是否有会议软件在运行"""
        try:
            import subprocess
            result = subprocess.run(
                ["tasklist.exe", "/FO", "CSV", "/NH"],
                capture_output=True, text=True, timeout=5
            )
            if result.returncode != 0:
                return None
            lines = result.stdout.strip().split('\n')
            running_processes = set()
            for line in lines:
                line = line.strip().strip('"')
                # CSV format: "name","pid","session","session#","mem"
                parts = line.split('","')
                if parts:
                    proc_name = parts[0].strip('"').lower()
                    running_processes.add(proc_name)

            for proc, name in self.MEETING_PROCESSES.items():
                if proc.lower() in running_processes:
                    return name
            return None
        except Exception as e:
            print(f"[MeetingMonitor] check error: {e}")
            return None

    def start_monitoring(self, on_start, on_stop, on_status=None):
        """启动监控"""
        if self._monitoring:
            return False
        self._monitoring = True
        self._on_start_callback = on_start
        self._on_stop_callback = on_stop
        self._on_status_callback = on_status
        self._thread = threading.Thread(target=self._monitor_loop, daemon=True)
        self._thread.start()
        return True

    def stop_monitoring(self):
        """停止监控"""
        self._monitoring = False
        if self._thread:
            self._thread.join(timeout=5)
            self._thread = None

    @property
    def is_monitoring(self):
        return self._monitoring

    @property
    def meeting_detected(self):
        return self._meeting_detected

    @property
    def detected_name(self):
        return self._detected_name

    def _monitor_loop(self):
        """监控循环"""
        if self._on_status_callback:
            self._on_status_callback("monitoring", "监控中 - 等待会议软件启动...")
        while self._monitoring:
            meeting_name = self._check_meeting_running()
            if meeting_name and not self._meeting_detected:
                # 会议软件刚启动
                self._meeting_detected = True
                self._detected_name = meeting_name
                if self._on_status_callback:
                    self._on_status_callback("detected", f"检测到{meeting_name}，3秒后自动录音...")
                # 等待3秒确认不是闪退
                time.sleep(3)
                if not self._monitoring:
                    break
                meeting_name2 = self._check_meeting_running()
                if meeting_name2:
                    if self._on_start_callback:
                        self._on_start_callback(meeting_name2)
                else:
                    self._meeting_detected = False
                    self._detected_name = ""
                    if self._on_status_callback:
                        self._on_status_callback("monitoring", "监控中 - 等待会议软件启动...")
                    continue
            elif not meeting_name and self._meeting_detected:
                # 会议软件已退出
                self._meeting_detected = False
                name = self._detected_name
                self._detected_name = ""
                if self._on_stop_callback:
                    self._on_stop_callback(name)
                if self._on_status_callback:
                    self._on_status_callback("monitoring", f"{name}已退出，继续监控...")
            elif meeting_name and self._meeting_detected:
                if self._on_status_callback:
                    self._on_status_callback("recording", f"监控中 - {meeting_name}运行中")
            time.sleep(self._check_interval)


# ─── 全局状态 ───
class AppState:
    def __init__(self):
        self.is_recording = False
        self.is_realtime = False
        self.recording_start = None
        self.transcript_text = ""
        self.sentence_info = []     # 说话人分离结果
        self.speaker_count = 0      # 识别到的说话人数量
        self.config = load_config()
        self.event_id = 0
        self.websocket = None
        self.event_loop = None  # uvicorn's event loop, set on first WS connect
        self.recorder = AudioRecorder()
        self.funasr = FunASRTranscriber()
        self.xfyun = None
        self.monitor = MeetingMonitor()
        self._lock = threading.Lock()
        self._realtime_thread = None

    def append_transcript(self, text):
        """Thread-safe append to transcript_text"""
        with self._lock:
            self.transcript_text += text
            return self.transcript_text

    def set_transcript(self, text):
        """Thread-safe set transcript_text"""
        with self._lock:
            self.transcript_text = text

    def next_event_id(self):
        self.event_id += 1
        return self.event_id

    def push_event_sync(self, ws, event_type, data=None):
        """Async push (call from async context like WS handler)"""
        event = {"id": self.next_event_id(), "type": event_type}
        if data is not None:
            event["data"] = data
        return ws.send_json(event)

    def push_from_thread(self, event_type, data=None):
        """Thread-safe push from background threads"""
        import asyncio
        loop = self.event_loop
        ws = self.websocket
        if not loop or not ws:
            print(f"[WARN] Cannot push {event_type}: no loop or ws")
            return
        event = {"id": self.next_event_id(), "type": event_type}
        if data is not None:
            event["data"] = data
        try:
            asyncio.run_coroutine_threadsafe(ws.send_json(event), loop)
        except Exception as e:
            print(f"[WARN] push_from_thread error: {e}")

state = AppState()

# ─── FastAPI 后端 ───
def create_app():
    from fastapi import FastAPI, WebSocket, WebSocketDisconnect
    from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
    from fastapi.staticfiles import StaticFiles

    app = FastAPI(title=APP_NAME)

    # Serve UI
    @app.get("/")
    async def index():
        html_path = UI_DIR / "index.html"
        return HTMLResponse(html_path.read_text(encoding="utf-8"))

    # REST: get config
    @app.get("/api/config")
    async def get_config():
        return JSONResponse(state.config)

    # REST: list audio devices
    @app.get("/api/devices")
    async def list_devices():
        try:
            devices = state.recorder.list_devices()
            return JSONResponse(devices)
        except Exception as e:
            return JSONResponse({"error": str(e)}, status_code=500)

    # WebSocket for real-time communication
    @app.websocket("/ws")
    async def websocket_endpoint(ws: WebSocket):
        await ws.accept()
        state.websocket = ws
        # Save uvicorn's event loop for background thread pushes
        import asyncio
        state.event_loop = asyncio.get_event_loop()
        # Send initial state
        await state.push_event_sync(ws, "config", state.config)
        # If there's existing transcript text, send it
        if state.transcript_text:
            await state.push_event_sync(ws, "transcript_ready", state.transcript_text)
        if state.is_recording:
            await state.push_event_sync(ws, "recording_changed", {"is_recording": True})
        if state.is_realtime:
            await state.push_event_sync(ws, "realtime_status", {"status": "recording", "message": "实时转写中"})
        await state.push_event_sync(ws, "status", "就绪")
        try:
            while True:
                msg = await ws.receive_json()
                await handle_message(ws, msg)
        except WebSocketDisconnect:
            state.websocket = None
        except Exception as e:
            print(f"WS error: {e}")
            traceback.print_exc()
            state.websocket = None

    async def handle_message(ws, msg):
        """Handle messages from frontend"""
        action = msg.get("action", "")

        if action == "get_config":
            await state.push_event_sync(ws, "config", state.config)

        elif action == "set_config":
            state.config.update(msg.get("data", {}))
            save_config(state.config)
            await state.push_event_sync(ws, "config_saved", {"ok": True})

        elif action == "list_devices":
            try:
                devices = state.recorder.list_devices()
                await state.push_event_sync(ws, "devices", devices)
            except Exception as e:
                await state.push_event_sync(ws, "log", {"message": f"获取设备列表失败: {e}"})

        elif action == "start_recording":
            source = msg.get("data", {}).get("source", state.config.get("audio_source", "mic"))
            try:
                ok = state.recorder.start(source=source)
                if ok:
                    state.is_recording = True
                    state.recording_start = time.time()
                    await state.push_event_sync(ws, "recording_changed", {"is_recording": True})
                    await state.push_event_sync(ws, "status", "录音中")
                    await state.push_event_sync(ws, "log", {"message": f"开始录音（音源: {source}）"})
                else:
                    await state.push_event_sync(ws, "log", {"message": "录音启动失败"})
            except Exception as e:
                await state.push_event_sync(ws, "log", {"message": f"录音启动失败: {str(e)}"})
                await state.push_event_sync(ws, "status", "就绪")

        elif action == "stop_recording":
            filepath = None
            try:
                filepath = state.recorder.stop()
            except Exception as e:
                await state.push_event_sync(ws, "log", {"message": f"停止录音失败: {str(e)}"})
            finally:
                # 必须无条件复位：旧实现把这两行放在 stop() 之后的 try 体内，
                # recorder.stop() 一抛异常就全部跳过，而 except 分支也不复位
                # → is_recording 永久卡在 True，前端按钮死在「录音中」，只能重启进程。
                # 录音确实已停（流已关或已不可用），状态就必须跟上。
                state.is_recording = False
                state.recording_start = None
                await state.push_event_sync(ws, "recording_changed", {"is_recording": False})
                await state.push_event_sync(ws, "status", "就绪")

            if filepath:
                await state.push_event_sync(ws, "log", {"message": f"录音已保存: {filepath}"})
                # 自动转写
                if state.config.get("auto_transcribe", True):
                    await state.push_event_sync(ws, "status", "转写中")
                    await state.push_event_sync(ws, "log", {"message": "开始转写录音文件..."})
                    # 在后台线程转写
                    threading.Thread(
                        target=_transcribe_file_task,
                        args=(filepath, ws),
                        daemon=True
                    ).start()
            else:
                await state.push_event_sync(ws, "log", {"message": "没有录到音频数据"})

        elif action == "start_realtime":
            if state.is_realtime:
                await state.push_event_sync(ws, "log", {"message": "实时转写已在运行"})
                return
            # 防止上一次转写线程还在清理中就启动新的
            if state._realtime_thread and state._realtime_thread.is_alive():
                await state.push_event_sync(ws, "log", {"message": "上一次转写仍在清理中，请稍候再试"})
                return
            state.is_realtime = True
            engine = state.config.get("engine", "FunASR")
            await state.push_event_sync(ws, "realtime_status", {"status": "loading", "message": f"正在启动{engine}..."})
            await state.push_event_sync(ws, "status", "加载中")
            # 启动实时转写线程
            state._realtime_thread = threading.Thread(
                target=_realtime_transcribe_task,
                args=(ws, engine),
                daemon=True
            )
            state._realtime_thread.start()

        elif action == "stop_realtime":
            state.is_realtime = False
            # 注意：recorder.stop() 由 _realtime_transcribe_task 线程负责调用，
            # 以便保存录音文件路径用于说话人分离
            if state.xfyun:
                result = state.xfyun.stop()
                if result:
                    state.append_transcript(result)
                state.xfyun = None  # 立即置空，防止 audio callback 访问已关闭的连接
            # 立即通知前端停止状态（不等后台线程清理完成）
            # 后台线程清理完成后会再次发送 stopped，前端幂等处理
            await state.push_event_sync(ws, "realtime_status", {"status": "stopped", "message": "实时转写已停止"})
            await state.push_event_sync(ws, "status", "就绪")
            await state.push_event_sync(ws, "log", {"message": "实时转写已停止"})
            # 后台线程会自行退出并清理（recorder.stop + 说话人分离）
            # 清理完成后会发送 transcript_ready 和最终的 realtime_status:stopped（幂等）
            # 不阻塞事件循环；start_realtime 会检查线程是否还在运行

        elif action == "transcribe_file":
            # 前端选择文件转写
            await state.push_event_sync(ws, "log", {"message": "请在设置中选择文件，或先录音再转写"})

        elif action == "generate_minutes":
            text = msg.get("data", {}).get("text", state.transcript_text)
            domain = msg.get("data", {}).get("domain", state.config.get("domain", "通用"))
            if not text.strip():
                await state.push_event_sync(ws, "log", {"message": "请先转写内容"})
                return
            await state.push_event_sync(ws, "status", "生成纪要中")
            # 说话人分离结果一并传入：带说话人/时间标签的转写对纪要质量帮助很大
            # （谁做的决策、谁认领的待办，纯文本里全丢了）
            threading.Thread(
                target=_generate_minutes_task,
                args=(text, domain, ws),
                kwargs={"sentence_info": list(state.sentence_info or [])},
                daemon=True
            ).start()

        elif action == "save_transcript":
            text = msg.get("data", {}).get("text", state.transcript_text)
            fmt = msg.get("data", {}).get("format", state.config.get("save_format", "txt"))
            filename = msg.get("data", {}).get("filename", f"transcript_{int(time.time())}")
            if not text.strip():
                await state.push_event_sync(ws, "log", {"message": "没有内容可保存"})
                return
            # 格式白名单：fmt 直接做文件后缀，不限定就能落盘 .bat/.ps1/.html
            if fmt not in ALLOWED_SAVE_FORMATS:
                await state.push_event_sync(ws, "log", {
                    "message": f"不支持的保存格式: {fmt}（可选 {'/'.join(ALLOWED_SAVE_FORMATS)}）"
                })
                return
            # 文件名来自前端，必须净化 + 归属校验，否则 "../../../evil"
            # 会直接写到 transcripts 目录外
            try:
                filepath = _resolve_output_path(TRANSCRIPTS_DIR, filename, fmt)
                filepath.parent.mkdir(parents=True, exist_ok=True)
            except Exception as e:
                await state.push_event_sync(ws, "log", {"message": f"保存路径无效: {e}"})
                return
            if fmt == "docx":
                # Word格式输出
                try:
                    # 计算录音时长：优先从sentence_info推算，否则从录音时间
                    recording_duration_s = None
                    if state.sentence_info:
                        max_end = max((s.get("end", 0) for s in state.sentence_info), default=0) / 1000
                        if max_end > 0:
                            recording_duration_s = max_end
                    _save_transcript_docx(
                        str(filepath), text, state.sentence_info, state.speaker_count,
                        recording_duration_s=recording_duration_s
                    )
                    await state.push_event_sync(ws, "log", {"message": f"已保存到 {filepath}"})
                except Exception as e:
                    await state.push_event_sync(ws, "log", {"message": f"Word保存失败: {str(e)}"})
            else:
                # 这里原本没有 try/except：磁盘满、权限不足、非法文件名都会抛
                # 异常，而它会一路冒泡到 websocket_endpoint 的兜底 except，
                # 把 state.websocket 置空 → UI 直接失联，用户只看到连接断。
                try:
                    filepath.write_text(text, encoding="utf-8")
                    await state.push_event_sync(ws, "log", {"message": f"已保存到 {filepath}"})
                except Exception as e:
                    await state.push_event_sync(ws, "log", {"message": f"保存失败: {str(e)}"})

        elif action == "copy_transcript":
            await state.push_event_sync(ws, "log", {"message": "已复制到剪贴板"})

        elif action == "start_monitor":
            if state.monitor.is_monitoring:
                await state.push_event_sync(ws, "log", {"message": "自动监控已在运行"})
                return

            def on_meeting_start(meeting_name):
                """检测到会议软件，自动开始录音"""
                source = state.config.get("audio_source", "system")
                try:
                    ok = state.recorder.start(source=source)
                    if ok:
                        state.is_recording = True
                        state.recording_start = time.time()
                        state.push_from_thread("recording_changed", {"is_recording": True})
                        state.push_from_thread("status", "录音中")
                        state.push_from_thread("log", {"message": f"检测到{meeting_name}，自动开始录音"})
                    else:
                        state.push_from_thread("log", {"message": "自动录音启动失败"})
                except Exception as e:
                    state.push_from_thread("log", {"message": f"自动录音失败: {str(e)}"})

            def on_meeting_stop(meeting_name):
                """会议软件退出，自动停止录音"""
                filepath = None
                try:
                    filepath = state.recorder.stop()
                except Exception as e:
                    state.push_from_thread("log", {"message": f"自动停止录音失败: {str(e)}"})
                finally:
                    # 同上：复位不能依赖 stop() 成功，否则自动监控路径下
                    # 也会把 is_recording 永久卡成 True
                    state.is_recording = False
                    state.recording_start = None
                    state.push_from_thread("recording_changed", {"is_recording": False})
                    state.push_from_thread("status", "就绪")

                if filepath:
                    state.push_from_thread("log", {"message": f"{meeting_name}已退出，录音已保存: {filepath}"})
                    if state.config.get("auto_transcribe", True):
                        state.push_from_thread("status", "转写中")
                        state.push_from_thread("log", {"message": "开始转写录音文件..."})
                        threading.Thread(
                            target=_transcribe_file_task,
                            args=(filepath, ws),
                            daemon=True
                        ).start()
                else:
                    state.push_from_thread("log", {"message": f"{meeting_name}已退出，但未录到音频"})

            def on_monitor_status(status, msg):
                state.push_from_thread("monitor_status", {"status": status, "message": msg})
                if status == "monitoring":
                    state.push_from_thread("status", "监控中")

            state.monitor.start_monitoring(on_meeting_start, on_meeting_stop, on_monitor_status)
            await state.push_event_sync(ws, "log", {"message": "自动监控已启动，等待会议软件..."})
            await state.push_event_sync(ws, "monitor_status", {"status": "monitoring", "message": "监控中 - 等待会议软件启动..."})

        elif action == "stop_monitor":
            if state.monitor.is_monitoring:
                state.monitor.stop_monitoring()
                await state.push_event_sync(ws, "log", {"message": "自动监控已停止"})
                await state.push_event_sync(ws, "monitor_status", {"status": "stopped"})
                if not state.is_recording:
                    await state.push_event_sync(ws, "status", "就绪")
            else:
                await state.push_event_sync(ws, "log", {"message": "自动监控未在运行"})

        else:
            await state.push_event_sync(ws, "log", {"message": f"未知操作: {action}"})

    return app


# ─── Word文档生成 ───
SPEAKER_COLORS = ["4472C4", "ED7D31", "70AD47", "FFC000", "5B9BD5", "A5A5A5", "7030A0", "C00000"]

def _format_timestamp(ms):
    """毫秒转 HH:MM:SS 格式"""
    s = ms / 1000
    h = int(s // 3600)
    m = int((s % 3600) // 60)
    sec = int(s % 60)
    return f"{h:02d}:{m:02d}:{sec:02d}"

def _format_timestamp_compact(ms):
    """毫秒转 MM:SS 或 HH:MM:SS（省略前导0小时）"""
    s = ms / 1000
    h = int(s // 3600)
    m = int((s % 3600) // 60)
    sec = int(s % 60)
    if h > 0:
        return f"{h:02d}:{m:02d}:{sec:02d}"
    return f"{m:02d}:{sec:02d}"

def _format_duration_cn(total_seconds):
    """秒转中文时长：'1时2分3秒' / '2分3秒' / '3秒'

    说话人发言时长动辄几十分钟，裸秒数（如 '2400.0秒'）无法快速判读。
    """
    if total_seconds is None:
        return ""
    total = int(total_seconds)
    if total < 0:
        total = 0
    h = total // 3600
    m = (total % 3600) // 60
    s = total % 60
    if h > 0:
        return f"{h}时{m}分{s}秒"
    if m > 0:
        return f"{m}分{s}秒"
    return f"{s}秒"

def _smart_paragraph_segment(sentence_info, gap_threshold=1500, long_pause_threshold=3000):
    """智能分段：将sentence_info合并为段落结构

    规则：
    1. 同一说话人连续句子，间隔 < gap_threshold 毫秒的合并为同一段
    2. 说话人切换时强制分段
    3. 停顿 > long_pause_threshold 毫秒作为分段点（即使同一说话人）
    4. 返回分段结构: [{spk, start, end, text, sentences: [{text, start, end}]}]
    """
    if not sentence_info:
        return []

    paragraphs = []
    current_para = None

    for s in sentence_info:
        spk = s.get("spk", 0)
        text = s.get("sentence", s.get("text", ""))
        start = s.get("start", 0)
        end = s.get("end", 0)

        should_new_para = False

        if current_para is None:
            should_new_para = True
        elif spk != current_para["spk"]:
            # 说话人切换 → 强制分段
            should_new_para = True
        elif start - current_para["end"] > long_pause_threshold:
            # 长停顿 → 分段
            should_new_para = True

        if should_new_para:
            if current_para is not None:
                paragraphs.append(current_para)
            current_para = {
                "spk": spk,
                "start": start,
                "end": end,
                "text": text,
                "sentences": [{"text": text, "start": start, "end": end}],
            }
        else:
            # 合并到当前段落
            current_para["end"] = end
            # 如果前文不以标点结尾且新文本不以标点开头，加逗号衔接
            if current_para["text"] and text:
                last_char = current_para["text"][-1]
                first_char = text[0]
                cn_puncts = set('。！？，、；：…—')
                if last_char not in cn_puncts and first_char not in cn_puncts:
                    current_para["text"] += "，" + text
                else:
                    current_para["text"] += text
            else:
                current_para["text"] += text
            current_para["sentences"].append({"text": text, "start": start, "end": end})

    if current_para is not None:
        paragraphs.append(current_para)

    return paragraphs

def _save_transcript_docx(filepath, text, sentence_info=None, speaker_count=0, recording_duration_s=None, timestamp_precision="sentence"):
    """生成Word格式的转写文档，支持说话人分离、智能分段和增强排版
    
    Args:
        filepath: 输出文件路径
        text: 纯文本转写结果
        sentence_info: 说话人分离结果 [{text, start, end, spk}, ...]
        speaker_count: 说话人数量
        recording_duration_s: 录音时长（秒），可选
        timestamp_precision: 时间戳精度 "sentence"(句级,默认) 或 "word"(词级)
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # ── 页面设置 A4 (210×297mm, 页边距2.5cm) ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 定义样式 ──
    # 标题样式
    style_title = doc.styles['Heading 1']
    style_title.font.size = Pt(22)
    style_title.font.name = "黑体"
    style_title.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    style_title.font.bold = True
    style_title.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    style_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_title.paragraph_format.space_after = Pt(12)

    # 正文样式
    style_normal = doc.styles['Normal']
    style_normal.font.size = Pt(12)
    style_normal.font.name = "宋体"
    style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)

    # ── 标题 ──
    title = doc.add_heading("会议转写记录", level=1)

    # ── 会议信息头 ──
    from datetime import datetime
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.autofit = True
    
    # 左列：转写时间
    cell_left = meta_table.rows[0].cells[0]
    cell_left.text = ""
    p_left = cell_left.paragraphs[0]
    run_label = p_left.add_run("转写时间：")
    run_label.font.size = Pt(10)
    run_label.font.name = "宋体"
    run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_label.font.bold = True
    run_label.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run_value = p_left.add_run(datetime.now().strftime('%Y年%m月%d日 %H:%M'))
    run_value.font.size = Pt(10)
    run_value.font.name = "宋体"
    run_value._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_value.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # 右列：录音时长 + 说话人数
    cell_right = meta_table.rows[0].cells[1]
    cell_right.text = ""
    p_right = cell_right.paragraphs[0]
    if recording_duration_s is not None:
        dur_str = _format_duration_cn(recording_duration_s)
        run_l1 = p_right.add_run("录音时长：")
        run_l1.font.size = Pt(10)
        run_l1.font.name = "宋体"
        run_l1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run_l1.font.bold = True
        run_l1.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run_v1 = p_right.add_run(dur_str)
        run_v1.font.size = Pt(10)
        run_v1.font.name = "宋体"
        run_v1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run_v1.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        if speaker_count > 0:
            run_sep = p_right.add_run("　")
            run_sep.font.size = Pt(10)
            run_l2 = p_right.add_run("说话人数：")
            run_l2.font.size = Pt(10)
            run_l2.font.name = "宋体"
            run_l2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run_l2.font.bold = True
            run_l2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run_v2 = p_right.add_run(f"{speaker_count}人")
            run_v2.font.size = Pt(10)
            run_v2.font.name = "宋体"
            run_v2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run_v2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    elif speaker_count > 0:
        run_l2 = p_right.add_run("说话人数：")
        run_l2.font.size = Pt(10)
        run_l2.font.name = "宋体"
        run_l2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run_l2.font.bold = True
        run_l2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run_v2 = p_right.add_run(f"{speaker_count}人")
        run_v2.font.size = Pt(10)
        run_v2.font.name = "宋体"
        run_v2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run_v2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # 去掉表格边框
    # 注意：OOXML 规定 w:tblPr 的子元素**必须有序**，w:tblBorders 要排在
    # w:tblLayout / w:tblLook 之前。直接 append 会得到
    # [tblW, tblLayout, tblLook, tblBorders] 这种非法顺序，Word 可能直接
    # 忽略该元素，甚至提示文档需要修复。改用 insert_element_before 按
    # schema 顺序插入。
    from docx.oxml import OxmlElement
    tblPr = meta_table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        borders.append(border)
    tblPr.insert_element_before(
        borders, 'w:shd', 'w:tblLayout', 'w:tblCellMar', 'w:tblLook'
    )

    doc.add_paragraph()  # 空行

    if sentence_info and len(sentence_info) > 0:
        # ── 说话人分离模式 ──
        
        # 说话人概览表
        if speaker_count > 1:
            # 说话人列表
            spk_stats = {}
            for s in sentence_info:
                spk = s.get("spk", 0)
                if spk not in spk_stats:
                    spk_stats[spk] = {"count": 0, "duration_ms": 0}
                spk_stats[spk]["count"] += 1
                start = s.get("start", 0)
                end = s.get("end", 0)
                spk_stats[spk]["duration_ms"] += (end - start)

            # 说话人概览标题
            overview_title = doc.add_paragraph()
            overview_title_run = overview_title.add_run("说话人概览")
            overview_title_run.font.size = Pt(14)
            overview_title_run.font.name = "黑体"
            overview_title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            overview_title_run.font.bold = True
            overview_title.paragraph_format.space_after = Pt(6)

            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            for i, h in enumerate(["说话人", "发言次数", "发言时长"]):
                hdr[i].text = h
                for p in hdr[i].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(10)
                        r.font.name = "宋体"
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            for spk_id in sorted(spk_stats.keys()):
                row = table.add_row().cells
                color_hex = SPEAKER_COLORS[spk_id % len(SPEAKER_COLORS)]
                row[0].text = f"说话人{spk_id + 1}"
                for p in row[0].paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(
                            int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
                        )
                        r.font.bold = True
                        r.font.size = Pt(10)
                        r.font.name = "宋体"
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                row[1].text = str(spk_stats[spk_id]["count"])
                for p in row[1].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
                        r.font.name = "宋体"
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                dur = spk_stats[spk_id]["duration_ms"] / 1000
                row[2].text = _format_duration_cn(dur)
                for p in row[2].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
                        r.font.name = "宋体"
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            doc.add_paragraph()  # 空行

        # ── 使用智能分段 ──
        paragraphs = _smart_paragraph_segment(sentence_info)
        
        # 转写内容标题
        content_title = doc.add_paragraph()
        content_title_run = content_title.add_run("转写内容")
        content_title_run.font.size = Pt(14)
        content_title_run.font.name = "黑体"
        content_title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        content_title_run.font.bold = True
        content_title.paragraph_format.space_after = Pt(6)

        prev_spk = None
        for para in paragraphs:
            spk = para["spk"]
            para_start = para["start"]
            para_end = para["end"]
            sentences = para["sentences"]
            color_hex = SPEAKER_COLORS[spk % len(SPEAKER_COLORS)]
            spk_color = RGBColor(
                int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
            )
            
            # 段落文本直接用 _smart_paragraph_segment 拼好的结果。
            # 不能在这里重新 join sentences：分段函数会在前句末尾无标点时
            # 补上“，”衔接，重新 join 会把这层处理全部丢掉，
            # 得到“今天讨论三件事第一是回款”这种粘连的句子。
            merged_text = para.get("text") or "".join(s["text"] for s in sentences)

            # 词级时间戳模式：逐句分行列出，与整段正文**二选一**。
            # 旧写法先输出整段正文、又附一行逐句带时间戳的文本，
            # 导致每句话在文档里**出现两次**（已实测确认）。
            per_sentence = timestamp_precision == "word" and len(sentences) > 1
            
            # 创建段落
            p = doc.add_paragraph()
            
            # 说话人切换时加大间距
            if spk != prev_spk:
                p.paragraph_format.space_before = Pt(12)
            else:
                p.paragraph_format.space_before = Pt(6)
            prev_spk = spk
            
            # 说话人标签（加粗+颜色）
            label_run = p.add_run(f"【说话人{spk + 1}】")
            label_run.font.bold = True
            label_run.font.size = Pt(12)
            label_run.font.name = "黑体"
            label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            label_run.font.color.rgb = spk_color
            
            # 时间段（如 00:05:23-00:08:45）
            time_range = f" {_format_timestamp(para_start)}-{_format_timestamp(para_end)}"
            time_run = p.add_run(time_range)
            time_run.font.size = Pt(9)
            time_run.font.name = "宋体"
            time_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            time_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            
            # 发言内容（首行缩进）。词级模式下改为逐句分行，不在此处输出整段。
            if not per_sentence:
                text_run = p.add_run(f"  {merged_text}")
                text_run.font.size = Pt(12)
                text_run.font.name = "宋体"
                text_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                p.paragraph_format.first_line_indent = Cm(0.74)  # 两字符缩进
            p.paragraph_format.space_after = Pt(6)

            # 词级时间戳模式：每句单独一行，行首带时间戳
            if per_sentence:
                for s in sentences:
                    sp = doc.add_paragraph()
                    sp.paragraph_format.space_before = Pt(0)
                    sp.paragraph_format.space_after = Pt(2)
                    sp.paragraph_format.left_indent = Cm(0.74)
                    ts_run = sp.add_run(f"[{_format_timestamp(s['start'])}] ")
                    ts_run.font.size = Pt(9)
                    ts_run.font.name = "宋体"
                    ts_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    ts_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                    body_run = sp.add_run(s["text"])
                    body_run.font.size = Pt(12)
                    body_run.font.name = "宋体"
                    body_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    else:
        # ── 普通模式：纯文本分段输出 ──
        paragraphs = re.split(r'([。！？\n])', text)
        combined = []
        buffer = ""
        for i, seg in enumerate(paragraphs):
            buffer += seg
            if seg in ('。', '！', '？', '\n') or i == len(paragraphs) - 1:
                if buffer.strip():
                    combined.append(buffer.strip())
                buffer = ""

        for para_text in combined:
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.size = Pt(12)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.paragraph_format.first_line_indent = Cm(0.74)  # 两字符缩进
            p.paragraph_format.space_after = Pt(6)

    doc.save(filepath)


def _transcribe_file_task(filepath, ws):
    """后台线程：转写录音文件"""
    try:
        engine = state.config.get("engine", "FunASR")
        if engine == "FunASR":
            def status_cb(status, msg):
                state.push_from_thread("realtime_status", {"status": status, "message": msg})

            speaker_diary = state.config.get("speaker_diarization", False)
            preset_spk = state.config.get("preset_spk_num", 0) or None
            hot_words = state.config.get("hot_words", "")

            if speaker_diary:
                # 说话人分离模式
                state.push_from_thread("log", {"message": "正在加载说话人分离模型..."})
                result = state.funasr.transcribe_file_with_diarization(
                    filepath, status_callback=status_cb, preset_spk_num=preset_spk,
                    hot_words=hot_words
                )
                if result is not None:
                    state.set_transcript(result.get("text", ""))
                    state.sentence_info = result.get("sentence_info", [])
                    state.speaker_count = result.get("speaker_count", 0)
                    # 发送结构化结果给前端
                    state.push_from_thread("transcript_ready", {
                        "text": state.transcript_text,
                        "sentence_info": state.sentence_info,
                        "speaker_count": state.speaker_count,
                    })
                    state.push_from_thread("status", "就绪")
                    spk_info = f"，识别{state.speaker_count}位说话人" if state.speaker_count > 0 else ""
                    state.push_from_thread("log", {"message": f"转写完成，共{len(state.transcript_text)}字{spk_info}"})
                else:
                    state.push_from_thread("status", "就绪")
                    state.push_from_thread("log", {"message": "说话人分离转写失败"})
            else:
                # 普通转写模式
                state.push_from_thread("log", {"message": "正在加载FunASR模型..."})
                file_model = state.config.get("funasr_model")
                result = state.funasr.transcribe_file(
                    filepath, status_callback=status_cb, hot_words=hot_words,
                    model_name=file_model,
                )
                if result is not None:
                    state.set_transcript(result)
                    state.sentence_info = []
                    state.speaker_count = 0
                    state.push_from_thread("transcript_ready", result)
                    state.push_from_thread("status", "就绪")
                    state.push_from_thread("log", {"message": f"转写完成，共{len(result)}字"})
                else:
                    state.push_from_thread("status", "就绪")
                    state.push_from_thread("log", {"message": "转写失败，未获得结果"})
        else:
            state.push_from_thread("log", {"message": "讯飞文件转写暂未实现，请使用FunASR引擎"})
            state.push_from_thread("status", "就绪")
    except Exception as e:
        print(f"[ERROR] _transcribe_file_task: {e}")
        traceback.print_exc()
        state.push_from_thread("log", {"message": f"转写失败: {str(e)}"})
        state.push_from_thread("status", "就绪")


def _realtime_transcribe_task(ws, engine="FunASR"):
    """后台线程：实时录音+转写"""
    def push(event_type, data=None):
        state.push_from_thread(event_type, data)

    try:
        # 启动录音
        source = state.config.get("audio_source", "mic")
        recording_filepath = None  # 录音文件路径（用于说话人分离）

        if engine == "xfyun":
            # 讯飞实时转写
            app_id = state.config.get("xfyun_app_id", "")
            api_key = state.config.get("xfyun_api_key", "")
            api_secret = state.config.get("xfyun_api_secret", "")
            if not all([app_id, api_key, api_secret]):
                push("realtime_status", {"status": "error", "message": "请先在设置中填写讯飞API密钥"})
                push("status", "就绪")
                state.is_realtime = False
                return

            state.xfyun = XfyunTranscriber(app_id, api_key, api_secret)

            def on_xfyun_result(text, is_end):
                state.append_transcript(text)
                push("transcript_partial", {"text": text, "is_end": is_end, "full_text": state.transcript_text})

            def on_xfyun_status(status, msg):
                push("realtime_status", {"status": status, "message": msg})
                if status == "ready":
                    push("status", "实时转写中")

            ok = state.xfyun.start(on_xfyun_result, on_xfyun_status)
            if not ok:
                push("realtime_status", {"status": "error", "message": "讯飞连接失败"})
                state.is_realtime = False
                return

            # 启动录音并转发音频
            def audio_to_xfyun(pcm_bytes):
                if state.xfyun and state.xfyun._connected:
                    state.xfyun.send_audio(pcm_bytes)

            try:
                state.recorder.start(source=source, callback=audio_to_xfyun)
            except Exception as e:
                push("realtime_status", {"status": "error", "message": f"录音启动失败: {str(e)}"})
                push("status", "就绪")
                state.is_realtime = False
                return
            push("realtime_status", {"status": "recording", "message": "讯飞实时转写中"})

            # 等待停止
            while state.is_realtime:
                time.sleep(0.5)

            # 停止录音，保存文件
            recording_filepath = state.recorder.stop()

        else:
            # FunASR 实时转写
            def funasr_status(status, msg):
                push("realtime_status", {"status": status, "message": msg})
                if status == "ready":
                    push("status", "实时转写中")

            if not state.funasr.load_model("iic/speech_paraformer-large_asr_nat-zh-cn-16k-common-vocab8404-online", status_callback=funasr_status):
                push("realtime_status", {"status": "error", "message": "FunASR模型加载失败"})
                push("status", "就绪")
                state.is_realtime = False
                return

            # 启动录音
            try:
                state.recorder.start(source=source)
            except Exception as e:
                push("realtime_status", {"status": "error", "message": f"录音启动失败: {str(e)}"})
                push("status", "就绪")
                state.is_realtime = False
                return
            push("realtime_status", {"status": "recording", "message": "FunASR实时转写中"})

            # 从队列取音频，送入FunASR
            chunk_size = [5, 10, 5]
            encoder_chunk_look_back = 4
            decoder_chunk_look_back = 1
            chunk_samples = chunk_size[1] * 160  # 1600 samples = 100ms

            cache = {}
            import numpy as np
            audio_buffer = np.array([], dtype=np.float32)

            # 标点恢复 + 智能分段状态
            raw_text = ""           # 原始无标点文本
            display_text = ""       # 带标点+分段的展示文本
            last_text_time = time.time()    # 上次收到新文字时间
            pause_threshold = 1.5   # 停顿超过1.5秒自动分段
            punc_running = False    # 标点线程是否正在运行

            # 实时热词：流式模型不吃 hotword 参数，只能在文本层纠
            hot_words = state.config.get("hot_words", "")
            hot_word_list = FunASRTranscriber._parse_hot_words(hot_words)
            if hot_word_list:
                if _POSTPROCESS_HOTWORDS_AVAILABLE:
                    push("log", {"message": f"实时热词纠正已启用（{len(hot_word_list)}个）"})
                else:
                    push("log", {"message": "已填热词但未安装 pypinyin/rapidfuzz，实时热词纠正已跳过"})

            # 定时标点恢复（异步线程，不阻塞主循环）
            # 必须定义在循环外：旧代码把它定在 `if audio_data is None:` 分支内，
            # 却在分支外调用 —— 若首次循环就取到音频（不进那个分支），
            # 会直接 NameError。
            def _async_punctuate(text_to_punctuate):
                nonlocal display_text, punc_running
                try:
                    punctuated = state.funasr.add_punctuation(text_to_punctuate)
                    # 热词纠正放在加标点之后：标点能隔开句子，避免跨句误匹配
                    if hot_word_list:
                        punctuated = correct_hotwords_by_pinyin(punctuated, hot_word_list)
                    # 停顿分段
                    if time.time() - last_text_time >= pause_threshold and not punctuated.endswith("\n"):
                        punctuated += "\n"
                    display_text = punctuated
                    push("transcript_partial", {"full_text": display_text})
                finally:
                    punc_running = False

            while state.is_realtime:
                audio_data = state.recorder.get_audio_chunk(timeout=0.5)
                if audio_data is None:
                    # 没有新数据，但buffer中有数据时也尝试处理
                    if len(audio_buffer) >= chunk_samples:
                        chunk = audio_buffer[:chunk_samples]
                        audio_buffer = audio_buffer[chunk_samples:]
                        try:
                            result = state.funasr.stream_model.generate(
                                input=chunk, cache=cache, is_final=False,
                                chunk_size=chunk_size,
                                encoder_chunk_look_back=encoder_chunk_look_back,
                                decoder_chunk_look_back=decoder_chunk_look_back,
                            )
                            if result and len(result) > 0:
                                text = result[0].get("text", "")
                                if text:
                                    raw_text += text
                                    last_text_time = time.time()
                        except Exception as e:
                            push("log", {"message": f"FunASR处理错误: {str(e)}"})
                    now = time.time()
                    if raw_text and now - last_text_time >= 1.5 and not punc_running:
                        punc_running = True
                        threading.Thread(target=_async_punctuate, args=(raw_text,), daemon=True).start()
                    continue

                try:
                    import numpy as np
                    audio_np = np.frombuffer(audio_data, dtype=np.int16).astype(np.float32) / 32767.0
                    audio_buffer = np.concatenate([audio_buffer, audio_np])

                    # 按chunk_samples大小逐块送入FunASR
                    while len(audio_buffer) >= chunk_samples:
                        chunk = audio_buffer[:chunk_samples]
                        audio_buffer = audio_buffer[chunk_samples:]

                        result = state.funasr.stream_model.generate(
                            input=chunk,
                            cache=cache,
                            is_final=False,
                            chunk_size=chunk_size,
                            encoder_chunk_look_back=encoder_chunk_look_back,
                            decoder_chunk_look_back=decoder_chunk_look_back,
                        )
                        if result and len(result) > 0:
                            text = result[0].get("text", "")
                            if text:
                                raw_text += text
                                last_text_time = time.time()
                except Exception as e:
                    push("log", {"message": f"FunASR处理错误: {str(e)}"})

                # 定时标点恢复（异步线程，不阻塞主循环）
                now = time.time()
                if raw_text and now - last_text_time >= 1.5 and not punc_running:
                    punc_running = True
                    threading.Thread(target=_async_punctuate, args=(raw_text,), daemon=True).start()

            # 最终flush - 处理buffer中剩余数据
            try:
                import numpy as np
                if len(audio_buffer) > 0:
                    if len(audio_buffer) < chunk_samples:
                        audio_buffer = np.concatenate([audio_buffer, np.zeros(chunk_samples - len(audio_buffer), dtype=np.float32)])
                    while len(audio_buffer) >= chunk_samples:
                        chunk = audio_buffer[:chunk_samples]
                        audio_buffer = audio_buffer[chunk_samples:]
                        result = state.funasr.stream_model.generate(
                            input=chunk, cache=cache, is_final=False,
                            chunk_size=chunk_size,
                            encoder_chunk_look_back=encoder_chunk_look_back,
                            decoder_chunk_look_back=decoder_chunk_look_back,
                        )
                        if result and len(result) > 0:
                            text = result[0].get("text", "")
                            if text:
                                raw_text += text
                # 发送结束标记
                result = state.funasr.stream_model.generate(
                    input=np.zeros(1, dtype=np.float32),
                    cache=cache,
                    is_final=True,
                    chunk_size=chunk_size,
                    encoder_chunk_look_back=encoder_chunk_look_back,
                    decoder_chunk_look_back=decoder_chunk_look_back,
                )
                if result and len(result) > 0:
                    text = result[0].get("text", "")
                    if text:
                        raw_text += text
            except:
                pass

            # 最终标点恢复 + 分段
            if raw_text:
                display_text = state.funasr.add_punctuation(raw_text)
                if hot_word_list:
                    display_text = correct_hotwords_by_pinyin(display_text, hot_word_list)
                state.set_transcript(display_text)
            else:
                state.set_transcript("")

            # 保存录音文件（用于说话人分离）
            recording_filepath = state.recorder.stop()

        # 完成 - 实时转写结束
        # 记录是否为用户主动停止（此时前端已收到stopped通知）
        was_user_stopped = not state.is_realtime

        # 如果开启了说话人分离，对录音文件做离线说话人分离
        speaker_diary = state.config.get("speaker_diarization", False)
        if speaker_diary and recording_filepath and engine == "FunASR":
            push("realtime_status", {"status": "diarizing", "message": "正在做说话人分离..."})
            push("status", "说话人分离中")
            push("log", {"message": "实时转写结束，正在对录音做说话人分离..."})
            try:
                preset_spk = state.config.get("preset_spk_num", 0) or None
                def diarization_status_cb(status, msg):
                    push("realtime_status", {"status": status, "message": msg})
                result = state.funasr.transcribe_file_with_diarization(
                    recording_filepath, status_callback=diarization_status_cb, preset_spk_num=preset_spk
                )
                if result is not None:
                    state.sentence_info = result.get("sentence_info", [])
                    state.speaker_count = result.get("speaker_count", 0)
                    # 发送结构化结果（含说话人信息）
                    push("transcript_ready", {
                        "text": state.transcript_text,
                        "sentence_info": state.sentence_info,
                        "speaker_count": state.speaker_count,
                    })
                    spk_info = f"，识别{state.speaker_count}位说话人" if state.speaker_count > 0 else ""
                    push("log", {"message": f"说话人分离完成{spk_info}"})
                else:
                    push("log", {"message": "说话人分离失败，仅输出纯文本转写结果"})
                    push("transcript_ready", state.transcript_text)
            except Exception as e:
                push("log", {"message": f"说话人分离出错: {str(e)}，仅输出纯文本转写结果"})
                push("transcript_ready", state.transcript_text)
        else:
            # 未开启说话人分离或无录音文件，发送纯文本结果
            push("transcript_ready", state.transcript_text)

        # 只有非用户主动停止时才发送stopped（用户主动停止时前端已收到通知）
        # 但说话人分离完成后需要再次发送stopped以恢复按钮状态
        if not was_user_stopped or speaker_diary:
            push("realtime_status", {"status": "stopped"})
        push("status", "就绪")

    except Exception as e:
        push("realtime_status", {"status": "error", "message": f"实时转写失败: {str(e)}"})
        push("status", "就绪")
        state.is_realtime = False
        traceback.print_exc()


# 实时流式模型（paraformer-online）不支持 hotword 参数，只能在文本层做热词纠正。
# 单字热词歧义太大（任意同音字都会命中），不参与拼音纠正。
_HOTWORD_MIN_LEN = 2
_CHAR_PINYIN_CACHE = {}


def _char_pinyin(ch):
    """单字拼音（带缓存）。

    不能直接对整段文本调 lazy_pinyin：它会把连续非汉字合并成一项
    （"abc中文" -> ['abc','zhong','wen']），导致字与音节不再一一对应，
    后续按下标切窗全会错位。改为逐字求拼音，结果缓存复用。
    """
    if ch in _CHAR_PINYIN_CACHE:
        return _CHAR_PINYIN_CACHE[ch]
    try:
        got = lazy_pinyin(ch)
        syl = (got[0] if got else ch).lower()
    except Exception:
        syl = ch.lower()
    _CHAR_PINYIN_CACHE[ch] = syl
    return syl


def _pinyin_syllables(s):
    """逐字拼音音节列表，与字符一一对应。"""
    return [_char_pinyin(c) for c in s]


# 不允许跨过这些字符做纠正（跨句纠正几乎必然是误匹配）
_HOTWORD_STOP_CHARS = set(" \t\r\n。！？，、；：…—“”‘’（）.!?,;:")


def correct_hotwords_by_pinyin(text, hot_words, threshold=0.85):
    """基于拼音相似度把转写文本里的近音错词纠正为热词。

    用于实时转写：paraformer-online 本身不接受 hotword 解码偏置，
    FunASR 的 postprocess_hotwords 也只在 AutoModel.generate 路径上生效，
    流式逐块 generate 拿不到。因此在文本层自己做一层。

    算法：建"音节 -> 出现位置"倒排，只在热词某个音节真实出现过的位置
    尝试开窗，避免对全文每个位置 x 每个热词暴力比对（长会议几万字会卡）。
    窗长取 len(w)、len(w)±1，容忍 ASR 多字/少字。命中后取不重叠区间，
    精确匹配优先于模糊匹配，长热词优先于短热词。

    依赖缺失（pypinyin/rapidfuzz）或无热词时原文返回，不报错。
    """
    words = FunASRTranscriber._parse_hot_words(hot_words)
    if not text or not words or not _POSTPROCESS_HOTWORDS_AVAILABLE:
        return text
    words = [w for w in words if len(w) >= _HOTWORD_MIN_LEN]
    if not words:
        return text

    text_syls = _pinyin_syllables(text)
    n = len(text)

    # 音节倒排：只在可能命中的位置尝试，把复杂度从 O(n×|words|) 降下来
    syl_positions = {}
    for idx, syl in enumerate(text_syls):
        syl_positions.setdefault(syl, []).append(idx)

    # 长热词优先，防止短热词先吃掉长热词的一部分
    words.sort(key=len, reverse=True)

    matches = []  # (start, end, word, score, word_len)
    seen_spans = set()
    for w in words:
        w_syls = _pinyin_syllables(w)
        w_key = "".join(w_syls)
        wl = len(w)
        starts = set()
        for k, syl in enumerate(w_syls):
            for p in syl_positions.get(syl, ()):
                s = p - k
                if 0 <= s < n:
                    starts.add(s)
        for s in sorted(starts):
            for span in (wl, wl + 1, wl - 1):
                if span < _HOTWORD_MIN_LEN or s + span > n:
                    continue
                key = (s, span, w)
                if key in seen_spans:
                    continue
                seen_spans.add(key)
                seg = text[s:s + span]
                if any(c in _HOTWORD_STOP_CHARS for c in seg):
                    continue
                if seg == w:
                    matches.append((s, s + span, w, 1.0, wl))
                    break
                score = _rf_fuzz.ratio("".join(text_syls[s:s + span]), w_key) / 100.0
                if score >= threshold:
                    matches.append((s, s + span, w, score, wl))
                    break

    if not matches:
        return text

    # 精确匹配 > 高分 > 长热词；同位置只采纳一个，区间不重叠
    matches.sort(key=lambda m: (m[0], -m[3], -m[4]))
    out = []
    cursor = 0
    for start, end, word, _score, _wl in matches:
        if start < cursor:
            continue
        out.append(text[cursor:start])
        out.append(word)
        cursor = end
    out.append(text[cursor:])
    return "".join(out)


# ─── 领域化处理 ───
# 旧实现里 domain 只是被塞进 prompt 的一个词（"领域：金融"），
# 模型无从知道该领域要重点抓什么、术语怎么写，等于没做领域化。
# 这里给每个领域补上关注点 + 易错术语提示。
DOMAIN_PROFILES = {
    "通用": {
        "focus": "决策结论、责任人、时限、未解决的分歧",
        "terms": "",
    },
    "金融": {
        "focus": "金额与币种、利率与费率、期限、授信与风控条件、合规要求、审批层级",
        "terms": "注意区分：本金/利息/罚息，授信/放款，年化/月化，基点(BP)与百分点",
    },
    "科技": {
        "focus": "技术方案取舍、架构与依赖、排期与里程碑、技术风险与兼容性",
        "terms": "英文缩写与产品名保持原文大小写（如 API、SDK、Kubernetes），不要译成中文",
    },
    "医疗": {
        "focus": "诊断结论、用药与剂量、检查检验指标、随访安排、知情同意与风险告知",
        "terms": "药名、剂量单位（mg/ml/IU）、指标数值必须与原文完全一致，绝不可推测或换算",
    },
    "教育": {
        "focus": "教学目标、课程与课时安排、考核评价方式、学生与家长诉求、师资安排",
        "terms": "注意区分学期/学年、必修/选修、学分与课时",
    },
    "法院": {
        "focus": "当事人主张、争议焦点、证据与举证责任、法律依据、程序节点与期限",
        "terms": "严格区分原告/被告/第三人，上诉/申诉/再审，事实认定与法律适用；"
                 "法条引用（如《民法典》第X条）须与原文一致",
    },
}


def _domain_guidance(domain):
    """生成领域化提示块。未知领域退化为通用，不抛异常。"""
    profile = DOMAIN_PROFILES.get(domain) or DOMAIN_PROFILES["通用"]
    lines = [f"【领域：{domain}】", f"本领域需重点抓取：{profile['focus']}"]
    if profile["terms"]:
        lines.append(f"术语规范：{profile['terms']}")
    return "\n".join(lines)


# 句末标点：分段时优先在这些位置切开，避免把一句话劈成两段
_SENTENCE_END_PUNCTS = "。！？；!?;\n"


def _split_text_for_llm(text, chunk_size, min_ratio=0.6):
    """按句子边界切分长文本，避免旧实现 text[i:i+CHUNK_SIZE] 的硬切。

    硬切会把一句话、一个数字、甚至一个说话人轮次劈到两段，
    两段各自摘要时都拿到残句 → 语义丢失、数据错配。

    策略：在 [chunk_size*min_ratio, chunk_size] 窗口内回溯找最后一个句末标点；
    找不到（如整段无标点）则退化为硬切，保证一定能推进不死循环。
    """
    if not text:
        return []
    if chunk_size <= 0:
        return [text]

    chunks = []
    start = 0
    total = len(text)
    min_len = max(1, int(chunk_size * min_ratio))

    while start < total:
        if total - start <= chunk_size:
            chunks.append(text[start:])
            break

        window_end = start + chunk_size
        cut = -1
        for i in range(window_end - 1, start + min_len - 1, -1):
            if text[i] in _SENTENCE_END_PUNCTS:
                cut = i + 1  # 标点归前一段
                break
        if cut <= start:
            cut = window_end  # 无标点可切，退化硬切
        chunks.append(text[start:cut])
        start = cut

    return [c for c in (c.strip() for c in chunks) if c]


def _build_annotated_transcript(sentence_info, max_chars=None):
    """把说话人分离结果渲染成带 [时间 说话人N] 标签的转写文本。

    说话人分离早就跑出了 spk 标签，但纪要 prompt 只吃纯文本，
    导致模型无法判断"谁说的" → 决策归属、待办负责人只能靠猜。
    复用 _smart_paragraph_segment 的分段结果，保证与 Word 导出一致。
    """
    if not sentence_info:
        return ""
    paragraphs = _smart_paragraph_segment(sentence_info)
    if not paragraphs:
        return ""

    lines = []
    used = 0
    for p in paragraphs:
        body = (p.get("text") or "").strip()
        if not body:
            continue
        line = f"[{_format_timestamp_compact(p.get('start', 0))} 说话人{p.get('spk', 0) + 1}] {body}"
        if max_chars is not None and used + len(line) > max_chars:
            break
        lines.append(line)
        used += len(line) + 1
    return "\n".join(lines)


def _generate_minutes_task(text, domain, ws, sentence_info=None):
    """后台线程：AI生成会议纪要（支持超长会议，分段摘要+汇总）"""
    def push(event_type, data=None):
        state.push_from_thread(event_type, data)

    def call_llm(prompt_text, cfg, llm_key, max_tokens=8192):
        """调用LLM API，返回文本结果"""
        body = cfg["body"]()
        messages = [{"role": "user", "content": prompt_text}]
        # 不同厂商的请求体层级不同（如 DashScope 将 messages 嵌在 input 下），
        # 由 cfg["inject"] 负责写到正确位置，不能一律写 body["messages"]
        cfg["inject"](body, messages, max_tokens)
        resp = requests.post(
            cfg["url"],
            headers=cfg["headers"](llm_key),
            json=body,
            timeout=cfg["timeout"],
        )
        if resp.status_code == 200:
            return cfg["extract"](resp.json())
        else:
            error_detail = ""
            try:
                payload = resp.json()
                # OpenAI 兼容格式：{"error": {"message": ...}}；DashScope：{"message": ...}
                error_detail = (
                    payload.get("error", {}).get("message", "")
                    or payload.get("message", "")
                )[:200]
            except Exception:
                error_detail = resp.text[:200]
            raise Exception(f"API返回{resp.status_code}: {error_detail}")

    try:
        llm_key = state.config.get("llm_api_key", "")
        if not llm_key:
            push("log", {"message": "请先在设置中填写AI纪要API Key"})
            push("status", "就绪")
            return

        provider = state.config.get("llm_provider", "xfyun")
        import requests

        # 将 messages / max_tokens 写入 OpenAI 兼容格式的请求体顶层
        def _inject_openai(body, messages, max_tokens):
            body["messages"] = messages
            body["max_tokens"] = max_tokens

        # DashScope 原生格式：messages 在 input 下，max_tokens 在 parameters 下
        def _inject_dashscope(body, messages, max_tokens):
            body.setdefault("input", {})["messages"] = messages
            body.setdefault("parameters", {})["max_tokens"] = max_tokens

        # API配置
        api_configs = {
            "xfyun": {
                "url": "https://maas-coding-api.cn-huabei-1.xf-yun.com/v2/chat/completions",
                "headers": lambda key: {"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
                "body": lambda: {"model": "astron-code-latest", "messages": [], "temperature": 0.3, "max_tokens": 8192},
                "inject": _inject_openai,
                "extract": lambda data: data["choices"][0]["message"]["content"],
                "name": "讯飞星火",
                "timeout": 180,
            },
            "volcengine": {
                "url": "https://ark.cn-beijing.volces.com/api/coding/v3/chat/completions",
                "headers": lambda key: {"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
                "body": lambda: {"model": "ark-code-latest", "messages": [], "temperature": 0.3, "max_tokens": 8192},
                "inject": _inject_openai,
                "extract": lambda data: data["choices"][0]["message"]["content"],
                "name": "火山引擎",
                "timeout": 180,
            },
            "deepseek": {
                "url": "https://api.deepseek.com/v1/chat/completions",
                "headers": lambda key: {"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
                "body": lambda: {"model": "deepseek-chat", "messages": [], "temperature": 0.3, "max_tokens": 8192},
                "inject": _inject_openai,
                "extract": lambda data: data["choices"][0]["message"]["content"],
                "name": "DeepSeek",
                "timeout": 120,
            },
            "qwen": {
                "url": "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation",
                "headers": lambda key: {"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
                "body": lambda: {
                    "model": "qwen-plus",
                    "input": {"messages": []},
                    "parameters": {"temperature": 0.3, "result_format": "message"},
                },
                "inject": _inject_dashscope,
                "extract": lambda data: data["output"]["choices"][0]["message"]["content"],
                "name": "通义千问",
                "timeout": 120,
            },
        }

        cfg = api_configs.get(provider, api_configs["xfyun"])

        # ── 说话人标注 ──
        # 有说话人分离结果时，用带 [时间 说话人N] 标签的文本喂模型，
        # 决策归属和待办负责人才有据可依；否则退回纯文本。
        annotated = _build_annotated_transcript(sentence_info)
        speaker_count = 0
        if sentence_info:
            speaker_count = max((s.get("spk", 0) for s in sentence_info), default=0) + 1
        if annotated:
            source_text = annotated
            push("log", {"message": f"已启用说话人标注纪要（{speaker_count}位说话人）"})
        else:
            source_text = text
        has_speakers = bool(annotated)

        # ── 分段策略 ──
        # 每段约6000字（留空间给prompt模板），4小时会议约3-4万字 → 5-7段
        CHUNK_SIZE = 6000
        text_len = len(source_text)
        domain_block = _domain_guidance(domain)

        if text_len <= CHUNK_SIZE:
            # 短会议：直接一次生成
            push("log", {"message": f"正在调用{cfg['name']}API生成纪要（全文{ text_len}字）..."})
            final_prompt = _build_final_prompt(source_text, domain, has_speakers=has_speakers)
            result = call_llm(final_prompt, cfg, llm_key)
            push("minutes_ready", {"text": result})
            push("status", "就绪")
            push("log", {"message": f"AI纪要生成完成（{cfg['name']}）"})
            return

        # 长会议：分段摘要 → 汇总
        # 按句末标点切，不再 text[i:i+CHUNK_SIZE] 硬切断句
        chunks = _split_text_for_llm(source_text, CHUNK_SIZE)

        push("log", {"message": f"长会议纪要：全文{text_len}字，分{len(chunks)}段处理..."})

        # 第一阶段：逐段生成摘要
        summaries = []
        speaker_hint = (
            "\n转写中的 [时间 说话人N] 是说话人标签，请据此判断观点、决策、待办分别属于谁，"
            "并在摘要中保留说话人标注。\n" if has_speakers else ""
        )
        for idx, chunk in enumerate(chunks):
            push("log", {"message": f"正在处理第{idx+1}/{len(chunks)}段..."})
            chunk_prompt = f"""你是一位资深会议秘书，请对以下会议转写片段提取关键信息。

{domain_block}
这是会议的第{idx+1}段（共{len(chunks)}段）。
{speaker_hint}
请提取：
1. 本段讨论的议题（如有新议题出现，明确标注）
2. 各方核心观点和争论焦点
3. 提及的关键数据（金额、比例、日期、人名等，务必与原文一致）
4. 达成的决策或共识
5. 待办事项及负责人

要求：
- 使用纯文本，禁止Markdown格式（不要用**、##、-等符号）
- 忠实于原文，不得编造
- 不确定的内容标注（待确认）
- 语言简洁正式

转写内容：
{chunk}"""
            try:
                summary = call_llm(chunk_prompt, cfg, llm_key, max_tokens=4096)
                summaries.append(summary)
            except Exception as e:
                push("log", {"message": f"第{idx+1}段处理失败: {str(e)[:60]}，跳过"})
                continue

        if not summaries:
            push("log", {"message": "所有分段处理失败，无法生成纪要"})
            push("status", "就绪")
            return

        # 第二阶段：汇总生成完整纪要
        push("log", {"message": f"分段摘要完成，正在汇总生成完整纪要..."})
        combined_summaries = "\n\n".join([f"【第{i+1}段摘要】\n{s}" for i, s in enumerate(summaries)])

        final_prompt = _build_final_prompt(
            combined_summaries, domain, is_summary=True, has_speakers=has_speakers
        )
        result = call_llm(final_prompt, cfg, llm_key, max_tokens=8192)

        push("minutes_ready", {"text": result})
        push("status", "就绪")
        push("log", {"message": f"AI纪要生成完成（{cfg['name']}，全文{text_len}字分{len(chunks)}段处理）"})

    except Exception as e:
        push("log", {"message": f"AI纪要生成失败: {str(e)}"})
        push("status", "就绪")


def _build_final_prompt(content, domain, is_summary=False, has_speakers=False):
    """构建最终纪要生成的prompt"""
    content_label = "各段摘要" if is_summary else "转写内容"
    speaker_block = ""
    if has_speakers:
        speaker_block = (
            "\n【说话人信息】\n"
            "内容中的 [时间 说话人N] 为自动说话人分离结果。请充分利用：\n"
            "- 参会人员按说话人编号列出；如转写中出现自我介绍或被称呼的姓名，则标注为“说话人1（张三）”\n"
            "- 各方观点需标明出自哪位说话人\n"
            "- 待办事项的负责人优先填说话人标识，不得根据猜测填写\n"
            "- 说话人编号与真实身份的对应关系不确定时，保留编号并注“（待确认）”\n"
        )
    return f"""你是一位资深会议秘书，请根据以下{content_label}撰写正式会议纪要。

{_domain_guidance(domain)}
{speaker_block}
【格式要求】
- 使用纯文本，禁止使用Markdown格式（不要用**、##、-等符号）
- 用中文数字编号（一、二、三...）标示大项，阿拉伯数字（1. 2. 3.）标示小项
- 段落之间空一行
- 关键数字、金额、日期用中文表述（如"三百二十万元"、"二〇二六年八月"）

【内容结构】
会议主题：（提炼会议核心议题）
会议时间：（如转写中提及）
参会人员：（如转写中提及，否则写"详见转写"）

一、会议背景
（简要说明召开此次会议的背景和目的，2-3句话）

二、主要议题与讨论
（按议题分条记录各方观点和讨论要点，每个议题3-5个要点，忠实反映不同意见）

三、关键决策
（明确记录达成的共识和决定，逐条列出）
四、待办事项
（需跟进的工作，格式：事项内容 → 负责人 → 完成时限，如未提及则写"无"）

五、风险与建议
（识别潜在风险，提出专业建议，如无明显风险则写"无"）

【注意事项】
- 忠实于转写内容，不得编造信息
- 转写中不确定的内容标注"（待确认）"
- 金额、比例等关键数据务必与原文一致
- 语言正式、简洁，避免口语化

{content_label}：
{content}"""


# ─── 启动 ───
def main():
    # 错误日志（打包后无控制台，错误写入文件便于排查）
    log_file = DATA_DIR / "startup_error.log"
    try:
        log_file.write_text("", encoding="utf-8")  # 清空旧日志
    except:
        pass

    def log_error(msg):
        import traceback
        ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        text = f"[{ts}] {msg}\n"
        try:
            log_file.write_text(text, encoding="utf-8")
        except:
            pass
        print(text)

    try:
        _main_inner(log_error)
    except Exception as e:
        import traceback
        log_error(f"启动失败:\n{traceback.format_exc()}")
        # 弹窗提示用户查看日志
        try:
            import ctypes
            ctypes.windll.user32.MessageBoxW(
                0, f"启动失败，请查看日志：\n{log_file}\n\n错误：{e}", "会议录音转写助手", 0x10
            )
        except:
            pass
        sys.exit(1)


def _main_inner(log_error):
    ensure_dirs()

    # Check if FastAPI is installed
    try:
        import fastapi
        import uvicorn
    except ImportError:
        print("安装依赖中...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "fastapi", "uvicorn", "websockets"])
        import fastapi
        import uvicorn

    app = create_app()

    # Find free port
    import socket
    port = None
    for p in range(18765, 18780):
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            s.bind(('127.0.0.1', p))
            s.close()
            break
        except OSError:
            continue
    else:
        p = 18765
    port = p

    # Start server in background
    def run_server():
        uvicorn.run(app, host="127.0.0.1", port=port, log_level="warning")

    server_thread = threading.Thread(target=run_server, daemon=True)
    server_thread.start()

    # Wait for server ready
    for _ in range(30):
        try:
            import urllib.request
            urllib.request.urlopen(f'http://127.0.0.1:{port}/', timeout=1)
            break
        except:
            time.sleep(0.3)

    # Pre-load FunASR model in background (saves 60+ seconds on first use)
    print("预加载语音模型...")
    threading.Thread(
        target=lambda: state.funasr.load_file_model(
            _resolve_file_model(state.config.get("funasr_model"))
        ),
        daemon=True,
    ).start()

    # Open Edge browser
    url = f"http://127.0.0.1:{port}/"
    print(f"启动 {APP_NAME} v{APP_VERSION}")
    print(f"访问地址: {url}")

    # Try Edge first, then default browser
    edge_paths = [
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
    ]
    browser_cmd = None
    for ep in edge_paths:
        if os.path.exists(ep):
            browser_cmd = ep
            break

    if browser_cmd:
        # Open in Edge app mode (no address bar, looks like native app)
        subprocess.Popen([
            browser_cmd,
            f"--app={url}",
            "--window-size=1400,860",
            f"--window-position=100,100",
        ])
    else:
        # Fallback to default browser
        subprocess.Popen(["cmd", "/c", "start", url])

    # Keep running
    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        print("\n退出")


if __name__ == "__main__":
    main()
