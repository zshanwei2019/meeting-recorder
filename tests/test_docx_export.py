"""Word 导出回归测试（直接 import app.py，不复制实现逻辑）。

覆盖 probe 实测确认的三处问题：
  1. 段落文本丢弃 _smart_paragraph_segment 的标点衔接 → 句子粘连
  2. timestamp_precision="word" 时正文与逐句时间戳双写 → 内容重复两遍
  3. w:tblBorders 直接 append 违反 OOXML 元素顺序

不加载 ASR 模型、不打开音频设备、不发网络请求。
运行：.venv\\Scripts\\python.exe tests\\test_docx_export.py
"""
import io
import os
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

import app  # noqa: E402
from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

PASS = []
FAIL = []


def check(name, got, expect):
    if got == expect:
        PASS.append(name)
        print(f"  OK   {name}")
    else:
        FAIL.append(name)
        print(f"  FAIL {name}\n         got={got!r}\n         expect={expect!r}")


def check_true(name, cond, hint=""):
    check(name + (f" ({hint})" if hint else ""), bool(cond), True)


def export(sentence_info, precision="sentence", speaker_count=2, dur=None):
    """跑一次真实导出，回传 (段落文本列表, 文件路径)。"""
    out = os.path.join(tempfile.mkdtemp(), "t.docx")
    app._save_transcript_docx(
        out,
        text="".join(s.get("sentence", "") for s in sentence_info),
        sentence_info=sentence_info,
        speaker_count=speaker_count,
        recording_duration_s=dur,
        timestamp_precision=precision,
    )
    doc = Document(out)
    return [p.text for p in doc.paragraphs if p.text.strip()], out


# 同一说话人两句、间隔小 → 会被合并进同一段，且前句末尾无标点
SAME_SPK = [
    {"sentence": "今天讨论三件事", "start": 0, "end": 1000, "spk": 0},
    {"sentence": "第一是回款", "start": 1200, "end": 2000, "spk": 0},
]

print("=== 1. 标点衔接不得丢失（旧 bug 复现）===")
# 旧写法 "".join(s["text"] for s in sentences) 会得到「三件事第一是」这种粘连。
# 分段函数本身补的逗号必须保留到 docx。
seg = app._smart_paragraph_segment(SAME_SPK)
check("分段函数确有补逗号", seg[0]["text"], "今天讨论三件事，第一是回款")
check("旧 join 写法确实丢逗号（证明修复必要）",
      "".join(s["text"] for s in seg[0]["sentences"]), "今天讨论三件事第一是回款")

texts, _ = export(SAME_SPK, "sentence")
body = "\n".join(texts)
check_true("导出正文保留逗号衔接", "今天讨论三件事，第一是回款" in body)
check_true("导出正文不含粘连写法", "三件事第一是" not in body)

print()
print("=== 2. word 模式内容不得重复（旧 bug 复现）===")
texts_w, _ = export(SAME_SPK, "word")
body_w = "\n".join(texts_w)
check("每句只出现一次", body_w.count("第一是回款"), 1)
check("首句也只出现一次", body_w.count("今天讨论三件事"), 1)
check_true("word 模式仍带时间戳", "[00:00:00]" in body_w)
check_true("每句独立成行",
           sum(1 for t in texts_w if t.strip().startswith("[")) == 2)

# sentence 模式不应出现逐句时间戳行
check_true("sentence 模式无逐句时间戳行",
           not any(t.strip().startswith("[") for t in texts))

print()
print("=== 3. 单句段落不受 word 模式影响 ===")
ONE = [{"sentence": "只有一句话", "start": 0, "end": 900, "spk": 0}]
t1, _ = export(ONE, "word")
b1 = "\n".join(t1)
check("单句仍只出现一次", b1.count("只有一句话"), 1)

print()
print("=== 4. tblBorders 必须符合 OOXML 元素顺序 ===")
# ECMA-376 CT_TblPrBase 规定的子元素顺序
ORDER = ["tblStyle", "tblpPr", "tblOverlap", "bidiVisual",
         "tblStyleRowBandSize", "tblStyleColBandSize",
         "tblW", "jc", "tblCellSpacing", "tblInd",
         "tblBorders", "shd", "tblLayout", "tblCellMar", "tblLook"]
idx = {n: i for i, n in enumerate(ORDER)}

_, path = export(SAME_SPK, "sentence", dur=3661)
doc = Document(path)
meta_tbl = doc.tables[0]._tbl
kids = [c.tag.split('}')[-1] for c in meta_tbl.tblPr]
pos = [idx[n] for n in kids if n in idx]
check_true("会议信息表子元素顺序合法", pos == sorted(pos), f"顺序={kids}")
check_true("tblBorders 确实写入了", "tblBorders" in kids)
check("tblBorders 不重复", kids.count("tblBorders"), 1)

print()
print("=== 5. 时长格式化 ===")
check("含时分秒", app._format_duration_cn(3661), "1时1分1秒")
check("仅分秒", app._format_duration_cn(125), "2分5秒")
check("仅秒", app._format_duration_cn(9), "9秒")
check("零秒", app._format_duration_cn(0), "0秒")
check("负数归零", app._format_duration_cn(-5), "0秒")
check("None 返回空串", app._format_duration_cn(None), "")
check("浮点截断", app._format_duration_cn(59.9), "59秒")
# 旧写法 f"{dur:.1f}秒" 对 40 分钟会给出 "2400.0秒"
check("40 分钟不再显示裸秒数", app._format_duration_cn(2400), "40分0秒")

print()
print("=== 6. 说话人概览表用可读时长 ===")
LONG = [
    {"sentence": "甲说了很久", "start": 0, "end": 2400000, "spk": 0},
    {"sentence": "乙简短回应", "start": 2400000, "end": 2405000, "spk": 1},
]
_, p2 = export(LONG, "sentence", speaker_count=2)
d2 = Document(p2)
overview = d2.tables[1]          # 0=会议信息 1=说话人概览
cells = [c.text for r in overview.rows for c in r.cells]
check_true("概览表含可读时长", "40分0秒" in cells, f"cells={cells}")
check_true("概览表不含裸秒数", "2400.0秒" not in cells)

print()
print("=== 7. 静态断言：防回归 ===")
src = io.open(ROOT / "app.py", encoding="utf-8").read()
check("不再用 append 挂 tblBorders", src.count("tblPr.append(borders)"), 0)
check_true("改用 insert_element_before", "insert_element_before" in src)
check("裸秒数格式已清除", src.count('f"{dur:.1f}秒"'), 0)
# 整段正文输出必须被 per_sentence 守卫，避免与逐句输出双写
i_guard = src.find("per_sentence = timestamp_precision ==")
i_body = src.find('text_run = p.add_run(f"  {merged_text}")')
check_true("per_sentence 守卫在正文输出之前", 0 < i_guard < i_body)
check_true("正文输出被 not per_sentence 包裹",
           "if not per_sentence:" in src)

print()
print(f"通过 {len(PASS)} 条，失败 {len(FAIL)} 条")
if FAIL:
    for f in FAIL:
        print(f"  - {f}")
sys.exit(1 if FAIL else 0)
